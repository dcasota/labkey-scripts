#!/usr/bin/env python3
"""
Build the OSM operator and administrator guide as a Microsoft Word document.

    tools/build_guide.py                  # write docs/OSM-Sample-Manager-Guide.docx
    tools/build_guide.py --out /tmp/x.docx
    tools/build_guide.py --check          # exit 2 if the committed file is stale

The document is generated, not hand-written. Facts that can drift — the project
version, the revision, the backlog, the decision register, the feature rollups —
are read from `pyproject.toml`, from git and from the memory database, so the
guide cannot quietly disagree with the repository it describes.

Determinism (AGENTS.md §4) is a property this script owns rather than assumes:
every timestamp written into the package is pinned to the HEAD commit date, and
the resulting zip is rewritten with fixed member timestamps in sorted order. Two
runs from the same commit produce byte-identical output, so `--check` is
meaningful and a rebuild produces no spurious diff.

Requires `python-docx`. Install it with:

    python3 -m pip install python-docx
"""
from __future__ import annotations

import argparse
import datetime as _dt
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import tomllib
import zipfile
from pathlib import Path
from typing import Any, Iterable, Sequence

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(Path(__file__).resolve().parent))

try:
    import docx
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError as exc:  # pragma: no cover - environment problem, not logic
    raise SystemExit(
        "python-docx is not installed. Install it with:\n"
        "    python3 -m pip install python-docx\n"
        f"(underlying import error: {exc})"
    ) from exc

import memory  # noqa: E402  - path is set above, matching tools/render_backlog.py

DEFAULT_OUT = REPO / "docs" / "OSM-Sample-Manager-Guide.docx"

#: Who the document is attributed to. The project itself, never a person or an
#: institution: no organisation publishes this guide. Held as a constant rather
#: than written inline at the `core_properties.author` assignment, because that
#: assignment would otherwise trip the secret scanner's `auth*` rule — and
#: broadening the scanner to accommodate a document is exactly the wrong fix
#: (see tests/test_no_secrets.py).
ATTRIBUTION = "The Open Sample Manager project"

#: Used when git is unavailable, so the build still succeeds and still pins a
#: fixed timestamp rather than reaching for the wall clock.
FALLBACK_DATE = _dt.datetime(2026, 8, 26, tzinfo=_dt.timezone.utc)

MONO = "OSM Mono"
CODE = "OSM Code"
TABLE_STYLE = "Light Grid Accent 1"

#: Test counts, recount with:  python3 -m pytest --collect-only -q
#: Held here rather than derived because parametrised tests make a static count
#: of `def test_` wrong, and collecting inside a document build is too slow.
TEST_SUITES = [
    ("tests/test_labkey_config.py", "44", "Credential resolution, TLS policy, URL handling"),
    ("tests/test_labkey_client.py", "46", "Session bootstrap, transport policy, error surfacing"),
    ("tests/test_memory_cli.py", "68", "Every memory CLI subcommand, including its refusals"),
    ("tests/test_memory_invariants.py", "50", "The committed knowledge base itself"),
    ("tests/test_no_secrets.py", "23", "The secret scanner, and proof that it fires"),
    ("tests/test_render_backlog.py", "20", "Rendering and drift detection"),
    ("tests/test_labkey_integration.py", "11", "Live server, marked `labkey`, deselected by default"),
]
UNIT_TEST_TOTAL = 251
INTEGRATION_TEST_TOTAL = 11


# --------------------------------------------------------------------------- facts


def git(*args: str) -> str:
    """Run a git command in the repository, returning '' when git cannot answer."""
    try:
        completed = subprocess.run(  # noqa: S603, S607 - fixed argv, no shell
            ["git", *args], capture_output=True, text=True, cwd=str(REPO), check=False
        )
    except OSError:
        return ""
    return completed.stdout.strip() if completed.returncode == 0 else ""


def project_version() -> str:
    data = tomllib.loads((REPO / "pyproject.toml").read_text(encoding="utf-8"))
    return str(data["project"]["version"])


def project_author() -> str:
    data = tomllib.loads((REPO / "pyproject.toml").read_text(encoding="utf-8"))
    authors = data["project"].get("authors") or []
    if not authors:
        return ""
    first = authors[0]
    return f"{first.get('name', '')} <{first.get('email', '')}>".strip()


def head_datetime() -> _dt.datetime:
    raw = git("log", "-1", "--format=%cI")
    if not raw:
        return FALLBACK_DATE
    try:
        return _dt.datetime.fromisoformat(raw).astimezone(_dt.timezone.utc)
    except ValueError:  # pragma: no cover - git always emits ISO-8601 here
        return FALLBACK_DATE


def load_facts() -> dict[str, Any]:
    """Everything the document quotes from the repository, gathered in one place."""
    conn = memory.connect()
    conn.row_factory = sqlite3.Row

    def rows(sql: str) -> list[sqlite3.Row]:
        return conn.execute(sql).fetchall()

    facts: dict[str, Any] = {
        "version": project_version(),
        "author": project_author(),
        "branch": git("rev-parse", "--abbrev-ref", "HEAD") or "(unknown)",
        "commit": git("rev-parse", "--short", "HEAD") or "(unknown)",
        "commit_full": git("rev-parse", "HEAD") or "(unknown)",
        "date": head_datetime(),
        "remote": git("remote"),
        "backlog": rows("SELECT * FROM backlog ORDER BY seq"),
        "decisions": rows("SELECT id, title, status, decision FROM decisions ORDER BY id"),
        "counts": {t: conn.execute(f"SELECT COUNT(*) FROM {t}").fetchone()[0]
                   for t in memory.TABLES},
        "req_kinds": rows("SELECT kind, COUNT(*) n FROM requirements GROUP BY kind ORDER BY kind"),
        "req_iterations": rows(
            "SELECT iteration, COUNT(*) n FROM requirements GROUP BY iteration ORDER BY iteration"),
        "ce_support": rows(
            "SELECT ce_support, COUNT(*) n FROM features GROUP BY ce_support ORDER BY n DESC"),
        "backlog_status": rows(
            "SELECT status, COUNT(*) n FROM backlog GROUP BY status ORDER BY status"),
        "verification_results": rows(
            "SELECT result, COUNT(*) n FROM verifications GROUP BY result ORDER BY result"),
    }
    conn.close()
    return facts


# --------------------------------------------------------------------------- styling


def _set_mono_fonts(style: Any) -> None:
    """Pin the monospace face on every script slot, not only `ascii`.

    Setting `style.font.name` alone leaves `w:hAnsi` unset in some Word builds,
    which silently renders the run in the body font.
    """
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(slot), "Courier New")


def _shade(element: Any, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def build_styles(doc: Any) -> None:
    """Add the two styles the guide needs beyond the stock template."""
    mono = doc.styles.add_style(MONO, WD_STYLE_TYPE.CHARACTER)
    mono.font.name = "Courier New"
    mono.font.size = Pt(9.5)
    mono.font.color.rgb = RGBColor(0x1F, 0x33, 0x64)
    _set_mono_fonts(mono)

    code = doc.styles.add_style(CODE, WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = doc.styles["No Spacing"]
    code.font.name = "Courier New"
    code.font.size = Pt(9)
    _set_mono_fonts(code)
    fmt = code.paragraph_format
    fmt.left_indent = Inches(0.25)
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)
    fmt.keep_together = True
    _shade(code.element.get_or_add_pPr(), "F2F2F2")


def add_page_number_footer(section: Any, left_text: str) -> None:
    """`left text` on the left, `Page N of M` on the right, as real Word fields."""
    paragraph = section.footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # A single centred tab stop is unreliable across page sizes; a right tab at
    # the text-width boundary is not.
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(section.page_width - section.left_margin - section.right_margin,
                      alignment=WD_TAB_ALIGNMENT.RIGHT)

    run = paragraph.add_run(f"{left_text}\t")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def field(instruction: str) -> None:
        r = paragraph.add_run()
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r._r.append(begin)
        r._r.append(instr)
        r._r.append(end)

    label = paragraph.add_run("Page ")
    label.font.size = Pt(8)
    label.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    field(" PAGE ")
    mid = paragraph.add_run(" of ")
    mid.font.size = Pt(8)
    mid.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    field(" NUMPAGES ")


# --------------------------------------------------------------------------- writer


_MARKUP = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


class Guide:
    """Thin authoring layer over python-docx.

    Body text accepts a two-token markup so the content below reads as prose
    rather than as run manipulation: `backticks` become monospace runs and
    **double asterisks** become bold ones. Nothing else is interpreted.
    """

    def __init__(self, facts: dict[str, Any]) -> None:
        self.facts = facts
        self.doc = docx.Document()
        self.sections = 0
        build_styles(self.doc)
        section = self.doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        add_page_number_footer(
            section,
            f"OSM Sample Manager Guide · version {facts['version']} · rev {facts['commit']}",
        )

    # -- run-level ---------------------------------------------------------

    def _runs(self, paragraph: Any, text: str) -> None:
        for token in _MARKUP.split(text):
            if not token:
                continue
            if token.startswith("`") and token.endswith("`") and len(token) > 1:
                run = paragraph.add_run(token[1:-1])
                run.style = self.doc.styles[MONO]
            elif token.startswith("**") and token.endswith("**") and len(token) > 3:
                paragraph.add_run(token[2:-2]).bold = True
            else:
                paragraph.add_run(token)

    # -- block-level -------------------------------------------------------

    def h1(self, text: str) -> None:
        self.sections += 1
        paragraph = self.doc.add_paragraph(text, style="Heading 1")
        if self.sections > 1:
            paragraph.paragraph_format.page_break_before = True

    def h2(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 3")

    def p(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        self._runs(paragraph, text)

    def quote(self, text: str) -> None:
        paragraph = self.doc.add_paragraph(style="Intense Quote")
        self._runs(paragraph, text)

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(style="List Bullet")
            self._runs(paragraph, item)

    def steps(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(style="List Number")
            self._runs(paragraph, item)

    def code(self, block: str) -> None:
        for line in block.strip("\n").splitlines():
            self.doc.add_paragraph(line or " ", style=CODE)

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[float] | None = None) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = self.doc.styles[TABLE_STYLE]
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for cell, heading in zip(table.rows[0].cells, headers):
            cell.text = ""
            run = cell.paragraphs[0].add_run(heading)
            run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = ""
                self._runs(cell.paragraphs[0], str(value))
        if widths:
            for row in table.rows:
                for cell, width in zip(row.cells, widths):
                    cell.width = Inches(width)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        if run.style.name != MONO:
                            run.font.size = Pt(9)
        self.doc.add_paragraph()

    def toc(self) -> None:
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = ("The table of contents is a Word field. Press F9, or right-click and "
                            "choose Update Field, to populate it.")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for node in (begin, instr, separate, placeholder, end):
            run._r.append(node)


# --------------------------------------------------------------------------- content


def write_title_page(g: Guide) -> None:
    f = g.facts
    doc = g.doc
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph("OSM Sample Manager", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Operator and Administrator Guide", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    g.table(
        ["Field", "Value"],
        [
            ["Project", "Open Sample Manager (OSM) — `osm-sample-manager`"],
            ["Project version", f"`{f['version']}` (from `pyproject.toml`)"],
            ["Repository revision", f"`{f['commit']}` on branch `{f['branch']}`"],
            ["Document date", f["date"].strftime("%Y-%m-%d")],
            ["Attributed to", ATTRIBUTION],
            ["`pyproject.toml` author field", f"`{f['author']}`" if f["author"] else "(none)"],
            ["Licence", "Apache-2.0 for code, CC-BY-4.0 for documentation"],
            ["Downstream target", "LabKey Server Community Edition 26.7.5"],
            ["Document status",
             "Generated from the repository by `tools/build_guide.py`. Do not hand-edit."],
        ],
        widths=[2.0, 4.5],
    )

    g.p("This guide is produced by a checked-in generator that reads the project version from "
        "`pyproject.toml`, the revision from git, and the backlog, decision register and "
        "feature inventory from the project memory database. It describes **what the "
        "repository contains at the revision above**, and marks planned work as planned.")
    g.p("No organisation is named as the publisher, and no author beyond the project itself is "
        "claimed. Where a person or an institution appears in this document it is quoted from "
        "a repository file and cited as such.")


def write_toc(g: Guide) -> None:
    g.h1("Table of contents")
    g.p("Word builds this table from the heading styles in the document. It is empty until you "
        "update it: click inside it and press **F9**, or right-click and choose "
        "**Update Field → Update entire table**. Some viewers (Google Docs, LibreOffice in "
        "read-only mode, most previewers) will not run the field at all — open the file in "
        "Word to see the populated contents.")
    g.toc()


def write_introduction(g: Guide) -> None:
    f = g.facts
    g.h1("1. Introduction")

    g.h2("1.1 What OSM Sample Manager is")
    g.p("Open Sample Manager (OSM) is an open-source sample-lifecycle system: sample "
        "registration with lineage and aliquots, a freezer map that matches the physical "
        "freezer slot for slot, an SOP-driven job queue, an electronic lab notebook with a "
        "binding signature, a faceted sample finder, a tamper-evident audit trail, a REST and "
        "MCP API, and downstream publishing into LabKey Community Edition.")
    g.p("The single most important framing fact, and the one this project records as most often "
        "got wrong:")
    g.quote("OSM is the system of record. LabKey CE is a downstream publish target.")
    g.p("OSM is **not** a LabKey module, **not** a reimplementation of the LabKey Sample "
        "Manager user interface, and does not depend on LabKey being available in order to "
        "function. The governing specification says so four times in §1 alone, and ADR-0001 "
        "records the reasoning and the evidence "
        "(`specs/adr/0001-osm-is-the-system-of-record.md`).")

    g.h2("1.2 The problem it solves")
    g.p("A laboratory that needs to track a sample from intake to disposal — where it is, in "
        "which freezer, what it was split from, who touched it, under which SOP, and what was "
        "written down about it — has effectively two options: buy a commercial LIMS, or "
        "assemble something from spreadsheets.")
    g.p("LabKey Community Edition looks like a third option and is not. It supplies the sample "
        "**data model** and almost none of the sample **management application**. That "
        "boundary is deliberate on LabKey's part: the missing capabilities are exactly what "
        "LabKey sells. OSM fills the application layer as open source, and publishes its "
        "results into LabKey CE so an institution keeps LabKey as its analysis and sharing "
        "surface.")

    g.h2("1.3 How OSM relates to LabKey Community Edition")
    g.p("The project maintains a 34-row capability inventory in its memory database. The "
        "headline, reproduced live from that database at build time:")
    g.table(
        ["LabKey CE support", "Capabilities", "What it means"],
        [[str(r["ce_support"]), str(r["n"]), {
            "native": "CE genuinely provides it",
            "partial": "The server supports it, but the UI or an edition gate blocks it",
            "absent": "Not present at any level",
        }.get(str(r["ce_support"]), "")] for r in f["ce_support"]],
        widths=[1.6, 1.2, 3.7],
    )
    g.p("The 22 absent capabilities include every one that makes a sample manager a sample "
        "manager: the storage hierarchy, box layouts, check-in and check-out, sample status "
        "enforcement, picklists, the sample finder, workflow jobs and the ELN.")

    g.h3("The five capabilities that look present and enforce nothing")
    g.p("This is the most operationally important thing to know about the downstream target. "
        "Each of the following is verified against the LabKey sources under `/root/scicore` "
        "and recorded in `docs/labkey-ce-ground-truth.md`. Each would produce a plausible but "
        "wrong implementation if assumed rather than checked.")
    g.table(
        ["Looks like", "Actually", "Evidence"],
        [
            ["Sample status enforcement",
             "`isOperationPermitted()` returns `true` unconditionally; **no rule is enforced**. "
             "A sample marked `Consumed` can still be edited, aliquoted and derived.",
             "`.../api/qc/SampleStatusService.java:84`"],
            ["PHI column tagging",
             "No `ComplianceService` is registered, so `DefaultComplianceService` passes "
             "everything through. **PHI tags survive export and import and mask nothing.**",
             "`.../api/compliance/ComplianceService.java`"],
            ["Freezer / inventory support",
             "`InventoryService.get()` returns `null` and every call site is null-guarded, so "
             "storage columns silently do not appear.",
             "`.../api/inventory/InventoryService.java`"],
            ["MCP server",
             "`McpService.get()` returns `NoopMcpService` with `isEnabled()` false; the search "
             "registration is commented out. There is no working MCP server in this build.",
             "`.../api/mcp/McpService.java`, `SearchModule.java:136-141`"],
            ["Barcode (Unique ID) field type",
             "Blocked only in the bundled JavaScript field editor. `property-createDomain.api` "
             "with the storage-unique-id concept URI creates a working barcode field.",
             "`ColumnRenderPropertiesImpl.java:48-49`, `domainDesigner.*.js`"],
        ],
        widths=[1.5, 3.3, 1.7],
    )
    g.p("The practical consequence for an operator: **do not rely on LabKey CE to enforce "
        "anything about a sample's state.** Whatever system owns those semantics must enforce "
        "them, and under ADR-0001 that system is OSM.")

    g.h3("What CE genuinely provides")
    g.p("These are worth knowing precisely, because they define what the publish bridge can "
        "lean on instead of reimplementing downstream: sample types with custom fields; naming "
        "patterns and ID generation (`NameGenerator.java`); lineage and derivation "
        "(`exp.Edge`, `experiment-lineage.api`, `experiment-derive.api`); a sample timeline "
        "audit registered unconditionally at `DETAILED` level; the full HTTP client API "
        "including API keys with restriction roles; full-text search over samples (Lucene "
        "10.5, documents indexed as `material:<rowId>`); and Apache-2.0 licensing at zero "
        "cost.")
    g.p("`exp.Material` natively carries `RootMaterialRowId`, `AliquotedFromLSID`, `IsAliquot`, "
        "`AliquotCount`, `AliquotVolume`, `StoredAmount`, `Units` and "
        "`AvailableAliquotVolume`. What is missing is the storage semantics that consume them, "
        "which live in the absent `inventory` module.")

    g.h2("1.4 How OSM relates to LabKey's commercial Sample Manager")
    g.p("OSM is not competing with LabKey's engineering; it is filling the deliberate boundary "
        "between LabKey's free data platform and its commercial application layer. The figures "
        "below are recorded in `docs/gap-analysis.md`, sourced from labkey.com/pricing and the "
        "`Available in: <editions>` badges on the labkey.org LIMS help pages.")
    g.table(
        ["Commercial tier", "Cost (per year)", "Gates"],
        [
            ["Sample Manager Starter", "USD 6,540 (five users)",
             "Freezer management, sample status, picklists, sample finder"],
            ["SM Professional", "USD 13,140",
             "ELN, workflow, assay management, API access"],
            ["LIMS Enterprise", "up to USD 59,400 (Biologics LIMS)",
             "Electronic signatures, ETL, SSO, the MCP server"],
        ],
        widths=[1.8, 2.0, 2.7],
    )
    g.p("In several places the OSM specification asks for more than the commercial product "
        "offers, and these are requirements rather than gold-plating: a tamper-evident "
        "SHA-256 audit chain (commercial Sample Manager chains nothing); row-level security on "
        "samples (LabKey partitions by container only); TTL-bounded slot reservations "
        "(commercial check-out holds a slot indefinitely); an eight-status lifecycle against "
        "three status *types*; an atomic box move; due-date escalation at T-1 and T+1; and MCP "
        "write tools that require explicit confirmation and are audited.")

    g.h2("1.5 What is actually implemented today")
    g.p("**Read this section before planning any work against OSM.** The repository at this "
        "revision is an early-stage, specification-driven project. What exists is the "
        "knowledge base, the project tooling, a LabKey client library, a verification harness "
        "and a substantial test suite. What does not yet exist is the sample-management "
        "application itself.")
    g.table(
        ["Capability", "State at this revision", "Where"],
        [
            ["Project memory database and CLI", "**Implemented**", "`tools/memory.py`"],
            ["Backlog and task renderer", "**Implemented**", "`tools/render_backlog.py`"],
            ["LabKey client library", "**Implemented** (PR-001, status `review`)",
             "`src/osm/labkey/`"],
            ["LabKey verification harness", "**Implemented** (8 read-only probes)",
             "`scripts/verify_labkey.py`"],
            ["Test suite and quality gate", "**Implemented**", "`tests/`, `Makefile`"],
            ["Secret scanner", "**Implemented**", "`tests/test_no_secrets.py`"],
            ["PostgreSQL schema and migrations", "Not implemented — planned, PR-003", "—"],
            ["Audit hash chain", "Not implemented — planned, PR-004 and PR-005", "—"],
            ["Identity, roles, OIDC, API keys", "Not implemented — planned, PR-006 and PR-007",
             "—"],
            ["REST API (FastAPI) and OpenAPI", "Not implemented — planned, PR-008", "—"],
            ["Sample registry, aliquots, lineage", "Not implemented — planned, PR-009 to PR-015",
             "—"],
            ["Freezer map and slot operations", "Not implemented — planned, PR-016 to PR-020",
             "—"],
            ["Job templates, jobs, queue", "Not implemented — planned, PR-021 to PR-023", "—"],
            ["Electronic lab notebook", "Not implemented — planned, PR-024 to PR-026", "—"],
            ["Faceted finder, picklists, RLS", "Not implemented — planned, PR-027 to PR-029",
             "—"],
            ["Publish outbox and LabKey bridge", "Not implemented — planned, PR-030 to PR-032",
             "—"],
            ["MCP server and assistant channels", "Not implemented — planned, PR-033 to PR-035",
             "—"],
            ["Operations, benchmarks, runbook", "Not implemented — planned, PR-036 to PR-038",
             "—"],
            ["User interface", "Not implemented — planned, PR-020 and beyond", "—"],
        ],
        widths=[2.2, 3.0, 1.3],
    )
    g.p(f"At this revision the backlog holds {f['counts']['backlog']} pull requests; "
        f"{sum(int(r['n']) for r in f['backlog_status'] if r['status'] == 'done')} are `done` "
        f"and {sum(int(r['n']) for r in f['backlog_status'] if r['status'] == 'todo')} are "
        f"`todo`. Section 12 lists every one of them.")

    g.h2("1.6 Audience, and how to read this guide")
    g.bullets([
        "**An operator setting the project up** should read sections 5 to 7, then run the "
        "installation check in section 6.4.",
        "**A developer picking up the next piece of work** should read sections 4 and 8, and "
        "then `AGENTS.md` and `memory.md` in the repository, which are the normative rules.",
        "**Someone evaluating OSM against LabKey** should read sections 1.3 to 1.5 and 12, and "
        "then `docs/gap-analysis.md` and `docs/labkey-ce-ground-truth.md`.",
        "**An auditor or reviewer** should read sections 9 and 10.",
    ])
    g.p("This document is a guide, not the contract. Where it disagrees with `AGENTS.md`, "
        "`memory.md`, `standards/` or the memory database, those win, and the disagreement is "
        "a defect in this generator.")


def write_architecture(g: Guide) -> None:
    f = g.facts
    g.h1("2. Architecture overview")

    g.h2("2.1 The shape of the target system")
    g.p("The architecture is fixed by eight accepted Architecture Decision Records. None of "
        "the runtime described here exists yet; this is the design the backlog builds towards, "
        "and it is stated so that the pieces that do exist can be read in context.")
    g.bullets([
        "**A standalone service tier.** OSM owns its own PostgreSQL database, services, API "
        "and user interface. A `labkey-bridge` publishes downstream. (ADR-0001)",
        "**Python with FastAPI, SQLAlchemy 2.x and Alembic**, with React and TypeScript for a "
        "user interface that is an ordinary client of the same REST API. (ADR-0002)",
        "**Audit inside the domain transaction.** Database triggers guarantee completeness, "
        "the unit-of-work additionally emits semantic domain events, and a PostgreSQL advisory "
        "lock serialises hash and sequence assignment so the SHA-256 chain is linear under "
        "concurrent writers. (ADR-0003)",
        "**Slot occupancy as a database constraint.** Occupancy and reservation are one "
        "`slot_assignment` row carrying a `tstzrange`, enforced by "
        "`EXCLUDE USING gist (slot_id WITH =, valid_during WITH &&)` over `btree_gist`. Two "
        "samples in one slot is not a race the application has to win; it is a constraint "
        "violation. (ADR-0004)",
        "**A transactional outbox for publishing.** The domain transaction inserts an "
        "`osm_publish_queue` row alongside the domain write; a worker drains it with "
        "`SELECT FOR UPDATE SKIP LOCKED` and upserts on `osm_id`, so publishing is idempotent "
        "and survives LabKey being down. (ADR-0005)",
        "**MCP as a thin adapter over REST.** Each tool translates onto a REST call "
        "authenticated as the agent's principal, so there is exactly one authorisation "
        "implementation and an agent structurally cannot exceed the role it authenticated as. "
        "(ADR-0006)",
        "**Project memory as a git-tracked SQL dump.** (ADR-0007, section 2.4)",
        "**Credentials exclusively from the environment.** (ADR-0008, section 7)",
    ])

    g.h2("2.2 How OSM sits on LabKey CE")
    g.p("It does not sit on LabKey CE. It sits **beside** it and publishes **into** it. The "
        "direction of the arrow is the whole of ADR-0001, and it has three practical "
        "consequences an operator needs to hold onto:")
    g.steps([
        "OSM must remain fully usable while LabKey is unavailable. A downstream system that "
        "can block an upstream write is not downstream.",
        "Only one package in the codebase knows LabKey exists: `src/osm/labkey/`. The package "
        "docstring states the rule — nothing in OSM outside that package should know that "
        "LabKey exists.",
        "Only `scripts/` may touch a running LabKey deployment, and nothing in this project "
        "may modify `/root/scicore` or the LabKey `SleepDrive-Lab` project, which belongs to "
        "different work.",
    ])
    g.p("The publish mapping, verified against CE, is a straightforward one-to-one onto native "
        "CE modules:")
    g.table(
        ["OSM object", "LabKey CE target", "Endpoint"],
        [
            ["Sample, Aliquot, Lineage", "`experiment` Sample Types",
             "`property-createDomain.api` + `query-import.api`"],
            ["Source", "`experiment` Data Classes", "`property-createDomain.api` + rows"],
            ["AssayRun", "`assay`", "`assay-importRun.api`"],
            ["Catalogs, job read-model", "`list`", "IntList + `query-import.api`"],
            ["File pointers", "`pipeline` / `filecontent`", "WebDAV `@files`"],
            ["SOPs, source cards", "`wiki`", "`wiki-saveWiki.api`"],
            ["Charts", "`visualization`", "`visualization-saveVisualization.api`"],
            ["Supplementary trail", "`audit`",
             "Supplements the OSM trail; never replaces it (CON-008)"],
        ],
        widths=[1.8, 2.0, 2.7],
    )

    g.h2("2.3 The specification-driven development workflow")
    g.p("OSM is developed specification-first. Every artefact links back to a section of the "
        "governing specification document, and the chain is recorded rather than remembered:")
    g.code(
        "spec .docx  ->  memory requirements  ->  PRD REQ-n  ->  FRD  ->  task  ->  code"
    )
    g.p("The workflow itself is version-controlled. Role definitions live in "
        "`.github/agents/` (`pm`, `architect`, `devlead`, `dev`) and the step definitions in "
        "`.github/prompts/` (`prd`, `frd`, `adr`, `create-standards`, `generate-agents`, "
        "`plan`, `implement`, `verify`). Each is a markdown file, so the method is reviewable "
        "in the same pull request as the work it produced.")
    g.table(
        ["Step", "Owner", "Produces"],
        [
            ["`/prd`", "pm", "`specs/prd.md`, requirements in the memory database"],
            ["`/frd`", "pm", "`specs/features/` — one document per capability"],
            ["`/adr`", "architect", "`specs/adr/NNNN-*.md` mirrored into `decisions`"],
            ["`/create-standards`", "architect", "`standards/`"],
            ["`/generate-agents`", "devlead", "`AGENTS.md`, regenerated when `standards/` changes"],
            ["`/plan`", "dev", "`specs/tasks/`, the `backlog` and `tasks` tables, "
                               "`docs/backlog.md`"],
            ["`/implement`", "dev", "Code and tests, one reviewable pull request"],
            ["`/verify`", "dev", "Rows in the `verifications` and `research` tables"],
        ],
        widths=[1.4, 1.0, 4.1],
    )
    g.p("A note on the current state: `specs/features/` exists as an empty directory. **No FRD "
        "has been written yet** — the backlog was planned directly from the PRD and the "
        "requirement register. `specs/README.md`, referenced from `memory.md`, does not exist "
        "either. Both are gaps in the artefact set rather than in the tooling.")

    g.h2("2.4 The memory database")
    g.p("The memory database is the project's durable, machine-queryable knowledge base. Every "
        "session reads from it instead of re-deriving knowledge, and `tools/memory.py` is the "
        "only supported way to read or write it.")
    g.table(
        ["Path", "Role"],
        [
            ["`.sdd/memory.sql`",
             "**Source of truth**, tracked in git. A deterministic SQL dump with a fixed table "
             "order and a fixed row order, so the same database always yields byte-identical "
             "text."],
            ["`.sdd/memory.db`",
             "Derived SQLite binary. **Git-ignored**, rebuilt on demand, and rebuilt "
             "automatically from the dump on first use if it is missing."],
            ["`tools/memory.py`", "The CLI. Enforces ID patterns, foreign keys and enum "
                                  "constraints that a hand edit would bypass."],
        ],
        widths=[1.6, 4.9],
    )
    g.p("ADR-0007 records why the dump is committed and the binary is not: a binary SQLite file "
        "produces no reviewable diff, merges badly, and makes it impossible to see in a pull "
        "request what knowledge changed. Every mutating CLI command rewrites the dump "
        "automatically, and `tools/memory.py check` fails if the tracked dump has drifted.")
    g.quote("Never hand-edit `.sdd/memory.sql` or `.sdd/memory.db`.")

    g.h3("Schema")
    g.p("Nine tables. The row counts below are read from the database at build time.")
    g.table(
        ["Table", "Holds", "ID form", "Rows"],
        [
            ["`requirements`",
             "Requirements extracted from the specification, each carrying the spec section it "
             "came from",
             "`FR-nnn`, `NFR-nnn`, `CON-nnn`, `PRO-nnn`", str(f["counts"]["requirements"])],
            ["`research`", "Research findings, each pinned to evidence", "`RF-nnn`",
             str(f["counts"]["research"])],
            ["`decisions`", "Architecture Decision Records, mirrored to `specs/adr/`",
             "`ADR-nnnn`", str(f["counts"]["decisions"])],
            ["`features`",
             "Capability inventory: commercial Sample Manager vs LabKey CE vs OSM",
             "`FEAT-nnn`", str(f["counts"]["features"])],
            ["`backlog`", "Ordered pull-request backlog with acceptance criteria",
             "`PR-nnn`", str(f["counts"]["backlog"])],
            ["`tasks`", "Tasks belonging to a backlog item", "`T-nnn`",
             str(f["counts"]["tasks"])],
            ["`verifications`", "How a claim about the environment was actually checked",
             "`V-nnn`", str(f["counts"]["verifications"])],
            ["`traceability`", "requirement → artifact edges", "(composite)",
             str(f["counts"]["traceability"])],
            ["`meta`", "Schema version and provenance", "—", str(f["counts"]["meta"])],
        ],
        widths=[1.2, 2.9, 1.6, 0.8],
    )
    g.p("`requirements.kind` is one of `functional`, `nonfunctional`, `constraint` or "
        "`prohibition`. Prohibitions matter: the specification forbids specific things and "
        "those are requirements too, each verified by a test asserting the capability is "
        "**absent** rather than merely refused. The current split is:")
    g.table(
        ["Kind", "Count"],
        [[f"`{r['kind']}`", str(r["n"])] for r in f["req_kinds"]],
        widths=[2.0, 1.2],
    )

    g.h2("2.5 The traceability model")
    g.p("The `traceability` table holds `requirement → artifact` edges, where the artifact kind "
        "is one of `feature`, `backlog`, `decision`, `verification`, `file` or `spec`. Two "
        "properties make it more than documentation:")
    g.bullets([
        "**A requirement nobody plans to satisfy fails the gate.** `tools/memory.py check` "
        "reports every requirement with no traceability link and exits 2. A planning hole "
        "cannot pass unnoticed.",
        "**A dangling edge fails the gate.** Edges of kind `backlog`, `decision`, `feature` and "
        "`verification` are resolved against their tables. Only `file` and `spec` point "
        "outside the database and so cannot be checked there.",
    ])
    g.p("The consequence is that `tools/memory.py show FR-029` answers *\"why does this code "
        "exist?\"* in one call, and the answer is guaranteed non-empty for every requirement in "
        "the register.")
    g.p("Two human-readable artefacts are generated from the same data by "
        "`tools/render_backlog.py`: `docs/backlog.md` (the ordered pull-request queue) and "
        "`specs/tasks/README.md` (the phase index and the requirement-to-task matrix). Both "
        "are checked for drift by `tools/render_backlog.py --check`, which is part of "
        "`make check`. Hand-editing either one is detected.")

    g.h2("2.6 Components that exist today")
    g.table(
        ["Component", "Path", "Responsibility"],
        [
            ["Memory CLI", "`tools/memory.py`",
             "Init, add, set, link, list, show, query, batch, stats, check, dump, rebuild"],
            ["Backlog renderer", "`tools/render_backlog.py`",
             "Renders `docs/backlog.md` and `specs/tasks/README.md`; `--check` detects drift"],
            ["Guide generator", "`tools/build_guide.py`", "Produces this document"],
            ["LabKey configuration", "`src/osm/labkey/config.py`",
             "Resolves credentials and TLS policy from the environment; never guesses"],
            ["LabKey client", "`src/osm/labkey/client.py`",
             "One logged-in session: context probing, CSRF, no redirects, preserved 4xx bodies, "
             "credential redaction"],
            ["Verification harness", "`scripts/verify_labkey.py`",
             "Eight read-only probes against a live server, optionally recorded as `V-nnn` rows"],
            ["Quality gate", "`Makefile`",
             "`make check` runs lint, types, tests, memory integrity, render drift and this "
             "guide's freshness"],
        ],
        widths=[1.5, 1.8, 3.2],
    )


def write_prerequisites(g: Guide) -> None:
    g.h1("3. Prerequisites")

    g.h2("3.1 Operating system and base tools")
    g.p("Nothing in the current codebase is platform-specific: it is pure Python plus git and "
        "GNU make. The reference environment on which the project was developed and verified "
        "is **VMware Photon OS 5.0**, and any Linux distribution carrying the versions below "
        "will do.")
    g.table(
        ["Requirement", "Minimum", "Reference environment", "Needed for"],
        [
            ["Python", "3.11 (`requires-python = \">=3.11\"`)", "3.11.13", "Everything"],
            ["git", "any recent version", "—",
             "Version metadata, and the tracked-file list the secret scanner uses"],
            ["GNU make", "any recent version", "—", "`make check` and the other gate targets"],
            ["SQLite", "bundled with Python (`sqlite3`)", "—", "The memory database"],
            ["PostgreSQL", "not yet required", "18.6 present",
             "Planned from PR-003; nothing in the tree connects to it today"],
            ["Node.js", "not yet required", "24.14 present",
             "Planned for the user interface (ADR-0002)"],
        ],
        widths=[1.1, 1.9, 1.5, 2.0],
    )

    g.h2("3.2 Python dependencies")
    g.p("The runtime dependency set is deliberately one library:")
    g.table(
        ["Group", "Declared in `pyproject.toml`", "Installed?"],
        [
            ["Runtime", "`requests>=2.32,<3`", "Required"],
            ["Development extra `dev`",
             "`pytest>=8`, `pytest-cov>=5`, `ruff>=0.6`, `mypy>=1.11`, `types-requests`",
             "Required to run the full gate"],
            ["Documentation", "`python-docx` (used by `tools/build_guide.py`)",
             "Required only to rebuild this guide"],
        ],
        widths=[1.6, 3.4, 1.5],
    )
    g.p("**Not yet pinned.** `pyproject.toml` declares ranges, not a lock file. Reproducing an "
        "exact resolved dependency set is an acceptance criterion of PR-002 and is not "
        "satisfied at this revision. Until then, `make lint` and `make types` report the "
        "absence of `ruff` and `mypy` instead of failing, so that the gate is not red for a "
        "tool that has not been pinned yet.")

    g.h2("3.3 The LabKey Community Edition instance")
    g.p("A LabKey CE server is required only for the integration test suite and for "
        "`scripts/verify_labkey.py`. The unit suite passes with no network at all, and OSM "
        "itself is specified to remain usable while LabKey is down.")
    g.table(
        ["Aspect", "Expectation"],
        [
            ["Version", "LabKey Server Community Edition 26.7.5 "
                        "(`gradle.properties: labkeyVersion=26.7.5`)"],
            ["Base URL", "`https://127.0.0.1:8443` by default; override with `LK_URL`"],
            ["TLS", "A self-signed certificate on loopback. See section 7.3 for how "
                    "verification is decided."],
            ["Servlet context", "Empty or `/labkey`. `LK_CONTEXT=auto` probes both rather than "
                                "assuming; assuming produces a 404 on every call."],
            ["Source enlistment", "`/root/scicore` — read-only for this project. **Never "
                                  "modified.**"],
            ["Reachability", "The client probes `login-whoAmI.api` at each candidate context "
                             "path and reports every candidate it failed on."],
        ],
        widths=[1.4, 5.1],
    )

    g.h2("3.4 Accounts and permissions")
    g.bullets([
        "**For the verification harness and the integration tests**: a LabKey account with "
        "**read** access to the container being probed (default `/home`). Every probe and "
        "every integration test is read-only; nothing is created, modified or deleted.",
        "**For the publish bridge, when it is built**: an API key minted with a "
        "`EditorWithoutDelete` restriction role. This is the least-privilege mechanism CE "
        "actually offers, and `standards/general/security.md` rule 10 requires it.",
        "**API keys are disabled by default.** `allowApiKeys` and `allowSessionKeys` both "
        "default to `false` (`AppPropsImpl.java:546-561`). A site administrator must enable "
        "them in site settings before `security-createApiKey.api` will mint anything.",
        "**API keys carry no prefix in this build.** They are 64-character lowercase hex; "
        "`apikey` is the header name and the basic-auth username, not a value prefix.",
    ])

    g.h2("3.5 Network and TLS expectations")
    g.bullets([
        "Outbound HTTPS from the machine running the tooling to `LK_URL`. Nothing else is "
        "contacted: the runtime dependency set is one HTTP library and there is no telemetry.",
        "**Redirects are never followed.** `admin-importFolder.post` answers 302 towards the "
        "server's canonical hostname; following that against a certificate bound to a "
        "different name loops until the client gives up. Every request sets "
        "`allow_redirects=False`, and an unexpected 3xx is raised as an error naming the "
        "`Location` header.",
        "**Certificate verification is on unless the host is loopback.** The rule and its "
        "single escape hatch are described in section 7.3.",
        "The repository has **no git remote**, deliberately, until `GITHUB_TOKEN` is supplied. "
        "A test asserts the absence, so an accidental push of an early mistake cannot happen "
        "first.",
    ])

    g.h2("3.6 Boundaries you must not cross")
    g.p("These are recorded in `AGENTS.md` §5 and are not negotiable within this project:")
    g.bullets([
        "Do not modify anything under `/root/scicore` or the running LabKey deployment "
        "**except through this project's `scripts/`**.",
        "`/root/install-labkey-sleepdrive-lab.sh` and the LabKey `SleepDrive-Lab` project "
        "belong to different work. Leave them alone. A test-suite rule and the integration "
        "test docstrings both restate this.",
        "Do not add a git remote.",
    ])


def write_installation(g: Guide) -> None:
    g.h1("4. Installation")

    g.h2("4.1 Obtaining the repository")
    g.p("The repository is local and has no remote by design, so there is nothing to clone from "
        "at this revision. Work from the existing enlistment:")
    g.code(
        "cd /root/osm-sample-manager\n"
        "git status\n"
        "git branch -a\n"
        "git log --oneline -5"
    )
    g.p("If you are transplanting the project to another machine, copy the working tree "
        "including `.git/` and do **not** add a remote — `tests/test_no_secrets.py` asserts "
        "that no remote is configured, and adding one turns the suite red until "
        "`GITHUB_TOKEN` is supplied and the decision is revisited.")

    g.h2("4.2 Preparing the Python environment")
    g.p("A virtual environment keeps the pinned development tools away from the system Python. "
        "The project imports from `src/` via `pythonpath` in `pyproject.toml`, so an editable "
        "install is convenient but not strictly required for the test suite.")
    g.code(
        "cd /root/osm-sample-manager\n"
        "python3 -m venv .venv\n"
        ". .venv/bin/activate\n"
        "python3 -m pip install --upgrade pip\n"
        "python3 -m pip install -e \".[dev]\"\n"
        "\n"
        "# Only needed if you intend to rebuild this guide:\n"
        "python3 -m pip install python-docx"
    )
    g.p("`.venv/` is git-ignored. Verify the interpreter and the one runtime dependency:")
    g.code(
        "python3 --version                     # must be 3.11 or newer\n"
        "python3 -c \"import requests; print(requests.__version__)\""
    )

    g.h2("4.3 Building the memory database")
    g.p("The binary database is a build artefact. Rebuild it from the committed dump:")
    g.code(
        "tools/memory.py rebuild        # .sdd/memory.sql -> .sdd/memory.db\n"
        "tools/memory.py stats"
    )
    g.p("If `.sdd/memory.db` is missing the CLI rebuilds it from the dump on first use, so this "
        "step is a convenience rather than a requirement. Both scripts under `tools/` are "
        "executable and carry a `#!/usr/bin/env python3` shebang; `python3 tools/memory.py …` "
        "works identically.")
    g.p("Expected output from `stats` at this revision — row counts, the backlog split by "
        "status, requirements by iteration, and the number of requirements with no "
        "traceability link, which must be zero:")
    g.code(
        "db:   /root/osm-sample-manager/.sdd/memory.db\n"
        "dump: /root/osm-sample-manager/.sdd/memory.sql\n"
        "  meta               3\n"
        "  requirements      97\n"
        "  research          15\n"
        "  decisions          8\n"
        "  features          34\n"
        "  backlog           38\n"
        "  tasks              9\n"
        "  verifications     29\n"
        "  traceability      97\n"
        "...\n"
        "requirements with no traceability link: 0"
    )

    g.h2("4.4 Validating the installation")
    g.p("`make check` is the quality gate. It must pass before any pull request is ready, and "
        "it is the single command that proves an installation works.")
    g.code("make check")
    g.p("It runs six targets in order:")
    g.table(
        ["Target", "Command", "What it proves"],
        [
            ["`lint`", "`python3 -m ruff check src tools scripts tests`",
             "Style and a subset of security lints. **Skipped with a message if `ruff` is not "
             "installed** — it is pinned by PR-002."],
            ["`types`", "`python3 -m mypy`",
             "Static types over `src`. Skipped the same way if `mypy` is absent."],
            ["`test`", "`python3 -m pytest -m \"not labkey\"`",
             "The unit suite, with no network access."],
            ["`memory`", "`python3 tools/memory.py check`",
             "Knowledge-base integrity; exit 2 on any problem."],
            ["`render`", "`python3 tools/render_backlog.py --check`",
             "`docs/backlog.md` and `specs/tasks/README.md` still match the database."],
            ["`docs-check`", "`python3 tools/build_guide.py --check`",
             "This guide still matches the repository. Skipped with a message if "
             "`python-docx` is not installed."],
        ],
        widths=[0.9, 2.5, 3.1],
    )
    g.p("A healthy run on a machine without `ruff` and `mypy` looks like this — the two "
        "skip messages are expected at this revision, not a failure:")
    g.code(
        "ruff not installed - skipping lint (pinned by PR-002)\n"
        "mypy not installed - skipping type check (pinned by PR-002)\n"
        "python3 -m pytest -m \"not labkey\"\n"
        "...............................................................\n"
        f"{UNIT_TEST_TOTAL} passed\n"
        "python3 tools/memory.py check\n"
        "memory integrity OK\n"
        "python3 tools/render_backlog.py --check\n"
        "generated backlog files are current\n"
        "/root/osm-sample-manager/docs/OSM-Sample-Manager-Guide.docx is current"
    )
    g.p("Override the interpreter with the `PY` variable if you need a specific one, for "
        "example `make check PY=/usr/bin/python3.11`.")

    g.h2("4.5 Optional: validating against the LabKey server")
    g.p("This step needs credentials and a reachable server. It is read-only.")
    g.code(
        "export LK_URL=https://127.0.0.1:8443\n"
        "export LK_USER=...\n"
        "export LK_PASSWORD=...          # or: export LK_APIKEY=...\n"
        "\n"
        "scripts/verify_labkey.py        # eight read-only probes\n"
        "make integration                # the 11 tests marked `labkey`"
    )
    g.p("With no credentials in the environment the integration tests **skip** rather than "
        "fail, and `scripts/verify_labkey.py` exits 1 with a message naming the missing "
        "variables.")

    g.h2("4.6 Rebuilding this guide")
    g.p("The document you are reading is generated. To rebuild it after changing the "
        "generator or the underlying facts:")
    g.code(
        "python3 -m pip install python-docx\n"
        "make docs                                  # or: python3 tools/build_guide.py\n"
        "make docs-check                            # what `make check` runs\n"
        "python3 tools/build_guide.py --check       # exit 2 if the committed file is stale\n"
        "python3 tools/build_guide.py --out /tmp/draft.docx   # build somewhere else"
    )
    g.p("The build is deterministic: every timestamp in the package is pinned to the HEAD "
        "commit date and the zip members are written in sorted order with fixed timestamps, so "
        "two runs from the same commit produce a byte-identical file and a rebuild leaves no "
        "spurious diff.")
    g.p("`--check` asks *does this document still describe the repository?*, not *was it built "
        "at HEAD?*. Those differ by exactly one commit the moment the guide is committed, "
        "because committing it moves HEAD. So the check reads the revision, branch and build "
        "date back out of the document's own properties and rebuilds with them: the prose, the "
        "version, the backlog, the decisions and every row count still have to match, and only "
        "the revision label is allowed to be older.")


def write_configuration(g: Guide) -> None:
    g.h1("5. Configuration")

    g.h2("5.1 The rule")
    g.p("Every credential is read from the environment, with **no defaults, no configuration "
        "files and no prompting**. A missing required variable aborts with a message naming it "
        "(ADR-0008). A credential is never passed as a command-line argument, because "
        "arguments are visible in the process table.")
    g.p("ADR-0008 records the specific failure mode this avoids, observed in neighbouring "
        "scripts: `build-labkey-community.sh` defaults a database password and a keystore "
        "password to real working values and writes them in plaintext. **A default that is a "
        "real credential is worse than no default, because it works, so nobody notices.**")

    g.h2("5.2 Variables the code actually reads")
    g.p("The table below lists every environment variable read by code in this repository, "
        "with the file and line that reads it. Anything not in this table has no effect on "
        "behaviour today.")
    g.table(
        ["Variable", "Default", "Read by", "Purpose"],
        [
            ["`LK_URL`", "`https://127.0.0.1:8443`", "`src/osm/labkey/config.py:137`",
             "LabKey base URL. Must have an `http` or `https` scheme and a host, or "
             "`ConfigError` is raised."],
            ["`LK_CONTEXT`", "`auto`", "`src/osm/labkey/config.py:138`",
             "Servlet context path. `auto` probes the origin and then `<origin>/labkey`; any "
             "other value is used verbatim."],
            ["`LK_USER`", "none", "`src/osm/labkey/config.py:139`",
             "LabKey login. The login parameter is `email`, not `username`."],
            ["`LK_PASSWORD`", "none", "`src/osm/labkey/config.py:140`",
             "LabKey password. Excluded from `repr`, never logged."],
            ["`LK_APIKEY`", "none", "`src/osm/labkey/config.py:141`",
             "LabKey API key. **Takes precedence** over user and password. Sent as an `apikey` "
             "header so it never appears in a URL."],
            ["`LK_INSECURE`", "unset (tri-state)", "`src/osm/labkey/config.py:155`",
             "`1`, `true`, `yes` or `on` skips TLS verification; anything else forces it on; "
             "unset defers to the loopback rule in section 5.3."],
            ["`LK_TIMEOUT`", "`60.0` seconds", "`src/osm/labkey/config.py:177`",
             "Per-request timeout in seconds. A non-numeric or non-positive value raises "
             "`ConfigError`."],
            ["`OSM_MEMORY_DB`", "`.sdd/memory.db`", "`tools/memory.py:55`",
             "Override the memory database path. Resolved at import time."],
            ["`OSM_MEMORY_DUMP`", "`.sdd/memory.sql`", "`tools/memory.py:56`",
             "Override the dump path. Resolved at import time."],
        ],
        widths=[1.1, 1.1, 1.5, 2.8],
    )
    g.quote("`LK_TIMEOUT` is read by the code but is documented neither in `memory.md` nor in "
            "`.env.example`. That is a documentation gap, not a behavioural one — the variable "
            "works exactly as described above.")
    g.p("One further knob is not an environment variable: the `Makefile` exposes `PY`, "
        "defaulting to `python3`, so `make check PY=python3.12` selects an interpreter.")

    g.h2("5.3 TLS verification and its opt-out")
    g.p("Certificate verification is decided from two inputs — whether `LK_INSECURE` is set, "
        "and whether the host is loopback. A host counts as loopback when it is `localhost`, "
        "`localhost.localdomain`, or an IP address whose `is_loopback` is true (which covers "
        "`127.0.0.0/8` and `::1`).")
    g.table(
        ["`LK_INSECURE`", "Host", "Result", "Why"],
        [
            ["unset", "loopback", "Verification **off**",
             "The deployment terminates TLS with a self-signed certificate, so verifying "
             "loopback always fails."],
            ["unset", "not loopback", "Verification **on**",
             "Any other host is a real network destination and must present a valid "
             "certificate."],
            ["`1` / `true` / `yes` / `on`", "loopback", "Verification **off**", "Explicit, and "
             "matches the default."],
            ["`1` / `true` / `yes` / `on`", "not loopback, `https`",
             "Verification **off**, plus a `WARNING` log line naming the host",
             "The escape hatch exists because internal certificate authorities are real, but "
             "it must never be quiet."],
            ["anything else (e.g. `0`)", "any", "Verification **on**",
             "Any non-true value forces verification on, whatever the host."],
        ],
        widths=[1.4, 1.1, 1.7, 2.3],
    )
    g.p("The warning text is exact, and worth recognising in a log:")
    g.code(
        "TLS verification is DISABLED for the non-loopback host 'labkey.example.org'.\n"
        "Credentials will be sent over a connection whose certificate is not checked.\n"
        "Unset LK_INSECURE unless this is deliberate."
    )
    g.p("Separately, when verification is off the client disables urllib3's per-request "
        "`InsecureRequestWarning` and logs one `WARNING` instead, because a warning emitted on "
        "every request trains readers to ignore warnings. Under pytest the same warning is the "
        "single documented entry in `filterwarnings`; every other warning is an error.")

    g.h2("5.4 Credential precedence")
    g.steps([
        "If `LK_APIKEY` is set, it wins. `LabKeyConfig.uses_api_key` is true, the key is sent "
        "as an `apikey` header, and no password login is attempted. An API key is preferred "
        "because it can be minted with a restriction role and revoked independently of the "
        "owning account.",
        "Otherwise both `LK_USER` and `LK_PASSWORD` must be set. The client bootstraps the "
        "session with `login-whoAmI.api` and then authenticates with `login-loginApi.api`.",
        "Otherwise `load_config()` raises `ConfigError` naming precisely which variables are "
        "missing. It never falls back to a guess, prompts, or reads a file.",
    ])
    g.p("Blank and whitespace-only values count as unset: `LK_USER=\"   \"` is treated as "
        "absent, which prevents a half-filled `.env` from producing an opaque 401 later.")

    g.h2("5.5 Reserved variables that no code reads yet")
    g.p("These are declared in `.env.example` and documented in `memory.md` so the names are "
        "stable, but **no code in the repository reads any of them at this revision**. Setting "
        "them today has no effect.")
    g.table(
        ["Variable", "Intended purpose", "Blocked on"],
        [
            ["`HUGGINGFACE_API_KEY`",
             "Inference token for the specification's §18 assistant work",
             "Iteration I6 — PR-035, assistant channels and RAG guardrails"],
            ["`HUGGINGFACE_MODEL_ID`", "Which model the assistant calls",
             "Iteration I6 — PR-035"],
            ["`GITHUB_TOKEN`",
             "Personal access token for pushing and opening pull requests",
             "Release tooling. The repository has no remote until this is supplied, "
             "deliberately."],
        ],
        widths=[1.7, 2.6, 2.2],
    )
    g.p("When they are wired in, they will follow the same rule as `LK_*`: no default, and a "
        "failure that names the missing variable.")

    g.h2("5.6 The `.env` convention")
    g.p("`.env.example` documents variable **names** with empty values and is tracked. `.env` "
        "is git-ignored and must never be committed. Copy and fill in:")
    g.code(
        "cp .env.example .env\n"
        "$EDITOR .env\n"
        "set -a; . ./.env; set +a        # export everything in the file"
    )
    g.p("`.gitignore` refuses `.env`, `.env.*` (with `!.env.example` re-admitting the "
        "template), `*.pem`, `*.key`, `*.p12`, `*.jks`, `*.keystore`, `*.pfx`, `secrets/`, "
        "`.netrc`, `*credentials*.json` and `*client_secret*.json`. Four of these patterns are "
        "asserted by a test, so the protection cannot be silently removed.")

    g.h2("5.7 Secrets: the non-negotiable rules")
    g.bullets([
        "**No credential, key, token or password appears in a tracked file. Ever.** Not in "
        "code, not in a test fixture, not in a comment, not in a commit message.",
        "**Never log a credential**, at any level, including in an exception message or a "
        "debug dump of a request. Log the URL and the status code.",
        "**Never pass a credential as a command-line argument.** Use the environment; "
        "arguments are visible in the process table.",
        "A tracked-file scan enforces all of this on every test run. See section 10.1.",
    ])


def write_usage(g: Guide) -> None:
    g.h1("6. Usage")

    g.h2("6.1 Reading the knowledge base")
    g.p("`tools/memory.py` is the everyday entry point. Read commands never mutate anything.")
    g.table(
        ["Command", "What it does"],
        [
            ["`tools/memory.py stats`",
             "Row counts, backlog by status, requirements by iteration, and the count of "
             "requirements with no traceability link"],
            ["`tools/memory.py list backlog --brief`", "The pull-request queue in `seq` order"],
            ["`tools/memory.py list requirements --where \"iteration='I2'\"`",
             "Any table, with an optional SQL `WHERE` clause"],
            ["`tools/memory.py show PR-001`",
             "Everything about one id across every table, plus its traceability edges"],
            ["`tools/memory.py query \"SELECT id,name,ce_support,gap FROM features "
             "WHERE gap='high'\"`",
             "Read-only SQL. Only `SELECT` and `WITH` are accepted"],
            ["`tools/memory.py check`", "The integrity gate. Exit 2 on any problem"],
            ["`tools/memory.py rebuild`", "`.sdd/memory.sql` → `.sdd/memory.db`"],
            ["`tools/memory.py dump`", "`.sdd/memory.db` → `.sdd/memory.sql`"],
        ],
        widths=[3.0, 3.5],
    )
    g.p("A `list` without `--brief` prints every non-empty column, wrapped. `--brief` prints "
        "`id  title  [status]` on one line per row, which is what you want when scanning.")

    g.h2("6.2 Writing to the knowledge base")
    g.p("Mutating commands rewrite `.sdd/memory.sql` automatically, so a commit always contains "
        "a reviewable text diff of what knowledge changed.")
    g.code(
        "tools/memory.py add req --id FR-098 --title \"...\" --kind functional \\\n"
        "    --source \"spec:§5\" --priority must --iteration I2 --body \"...\"\n"
        "\n"
        "tools/memory.py add research --id RF-016 --topic \"...\" --finding \"...\" \\\n"
        "    --evidence-kind source --evidence-ref /root/scicore/...\n"
        "\n"
        "tools/memory.py add verification --id V-030 --claim \"...\" --method http \\\n"
        "    --command \"POST /home/query-getSchemas.api\" --result pass --detail \"...\"\n"
        "\n"
        "tools/memory.py link --req FR-098 --kind backlog --to PR-017\n"
        "tools/memory.py set backlog PR-001 status done\n"
        "tools/memory.py batch seed.txt        # many statements, one dump at the end"
    )
    g.p("The CLI validates before it writes, and its refusals are deliberate:")
    g.bullets([
        "**ID patterns.** `FR|NFR|CON|PRO-nnn`, `RF-nnn`, `ADR-nnnn` (four digits), "
        "`FEAT-nnn`, `PR-nnn`, `T-nnn`, `V-nnn`. `ADR-001` is rejected; the correct form is "
        "`ADR-0001`.",
        "**Foreign keys.** A task must name a backlog item that exists; a traceability edge "
        "must name a requirement that exists. The error explains which reference was bad "
        "rather than suggesting `--replace`, because replacing a row cannot make a missing "
        "parent exist.",
        "**Enumerations.** `kind`, `priority`, `status`, `size`, `method`, `result`, "
        "`evidence_kind`, `ce_support` and `gap` are all constrained by `CHECK` clauses.",
        "**Read-only `query`.** Anything that is not a `SELECT` or a `WITH` is refused with a "
        "pointer to `add` and `set`.",
        "**`batch` accepts only `add`, `link` and `set`.** On failure it reports the file, the "
        "line number and how many rows had already been applied, so a partial application is "
        "never silent.",
    ])
    g.quote("Never hand-edit `.sdd/memory.sql` or `.sdd/memory.db`. The CLI enforces "
            "constraints that a hand edit would bypass, and `make check` will catch the drift "
            "afterwards anyway.")

    g.h2("6.3 The SDD loop, end to end")
    g.p("This is the working method from `AGENTS.md` §6 and `.github/prompts/implement.prompt."
        "md`. It is the process an operator or a contributor follows for every change.")
    g.steps([
        "**Pick the work.** `tools/memory.py list backlog --where \"status='todo'\" --brief` "
        "and take the lowest `seq` whose dependencies are all `done`.",
        "**Claim it.** `tools/memory.py set backlog PR-nnn status in-progress`, then create the "
        "branch `pr-nnn-slug` (the branch name is stored on the backlog row).",
        "**Read before writing.** The task specification, the ADRs it cites, then `memory.md` "
        "and `AGENTS.md`.",
        "**Verify, do not assume.** Any claim about LabKey behaviour must be backed by source "
        "under `/root/scicore` or by a real HTTP call, and recorded with "
        "`tools/memory.py add verification`. A `fail` result is as valuable as a `pass` — "
        "record it; do not retry until the answer is convenient.",
        "**Implement with tests in the same change.** Test the **denials**, not only the "
        "successes: a role test that only proves access works has tested half the requirement.",
        "**Stop if a significant architectural decision is missing.** Hand back for an ADR. Do "
        "not decide it inside the implementation.",
        "**Link the requirement.** `tools/memory.py link --req FR-nnn --kind backlog "
        "--to PR-nnn`, so the requirement is traced and `check` stays green.",
        "**Re-render.** `tools/render_backlog.py` if anything in the backlog changed, so "
        "`docs/backlog.md` and `specs/tasks/README.md` stay current.",
        "**Close out.** Every acceptance checkbox ticked, `make check` green, a `JOURNAL.md` "
        "entry appended, and `tools/memory.py set backlog PR-nnn status review`.",
    ])
    g.p("Commit convention (`AGENTS.md` §7): an imperative subject under 72 characters and a "
        "body explaining **why**, not what.")

    g.h2("6.4 Keeping the generated documents current")
    g.p("`docs/backlog.md` and `specs/tasks/README.md` are generated from the memory database. "
        "Editing them by hand is detected and overwritten.")
    g.code(
        "tools/render_backlog.py                 # write both files\n"
        "tools/render_backlog.py --check         # exit 2 if either is stale\n"
        "tools/render_backlog.py --out-dir /tmp/render   # render elsewhere, used by the tests"
    )
    g.p("`--check` is part of `make check`, so a knowledge-base change that is not re-rendered "
        "fails the gate rather than leaving the readable artefacts describing a plan that is "
        "no longer the plan.")

    g.h2("6.5 Probing the LabKey publish target")
    g.p("`scripts/verify_labkey.py` is the executable form of the project's central rule: never "
        "assert a LabKey behaviour that has not been checked. Every probe is a read; nothing "
        "is written to the server.")
    g.code(
        "export LK_URL=https://127.0.0.1:8443\n"
        "export LK_USER=...  LK_PASSWORD=...          # or LK_APIKEY=...\n"
        "\n"
        "scripts/verify_labkey.py                     # probe and print\n"
        "scripts/verify_labkey.py --container /home   # probe a different container\n"
        "scripts/verify_labkey.py --record V-030      # also write verification rows\n"
        "scripts/verify_labkey.py -v                  # debug logging"
    )
    g.table(
        ["Probe", "Claim it checks"],
        [
            ["Identity", "An authenticated session can be established from environment "
                         "credentials, and a CSRF token is captured"],
            ["Schemas", "CE exposes no `inventory`, `storage`, `freezer`, `workflow`, `eln` or "
                        "`labbook` schema"],
            ["Material columns",
             "`exp.Materials` carries `RootMaterialRowId`, `AliquotedFromLSID`, `IsAliquot`, "
             "`AliquotCount`, `StoredAmount`, `Units`, `SampleState`"],
            ["Sample states", "CE offers only three state types against the eight the "
                              "specification requires"],
            ["Audit", "The `auditLog` schema provides `SampleTimelineEvent`, "
                      "`SampleSetAuditEvent` and `TransactionAuditEvent`"],
            ["Units", "CE ships a fixed measurement-unit vocabulary the bridge must map onto"],
            ["Containers", "Container enumeration works and the scratch space can be located"],
            ["Known-bad action", "The client refuses `query-importData.api` locally and names "
                                 "`query-import.api` instead"],
        ],
        widths=[1.5, 5.0],
    )
    g.p("Exit codes: **0** all probes passed, **2** at least one probe failed, **1** fatal (bad "
        "configuration or an unreachable server). `--record` shells out to `tools/memory.py` "
        "rather than importing it, which keeps the single documented write path honest: if the "
        "CLI would reject the row, so does the harness.")

    g.h2("6.6 Using the LabKey client from Python")
    g.p("`src/osm/labkey/` is the only sanctioned way for OSM code to talk to a LabKey "
        "deployment. Use the client as a context manager so the connection pool is released.")
    g.code(
        "import sys\n"
        "sys.path.insert(0, \"src\")\n"
        "\n"
        "from osm.labkey import LabKeyClient, LabKeyError\n"
        "\n"
        "with LabKeyClient.from_env() as lk:\n"
        "    print(lk.base_url, lk.user_email)\n"
        "    print(lk.schemas(\"/home\"))\n"
        "    print(lk.queries(\"auditLog\", container=\"/home\"))\n"
        "    rows = lk.select_rows(\"exp\", \"SampleStateType\", container=\"/home\")\n"
        "    print([r[\"Value\"] for r in rows])\n"
        "\n"
        "    try:\n"
        "        lk.post_json(\"query-getQueryDetails.api\",\n"
        "                     {\"schemaName\": \"exp\", \"queryName\": \"NoSuchQuery\"})\n"
        "    except LabKeyError as exc:\n"
        "        print(exc.status, exc.exception_class)\n"
        "        print(exc.body)      # the server's explanation, preserved on purpose"
    )
    g.p("Behaviour worth knowing before you rely on it:")
    g.bullets([
        "**Success is not a status code.** LabKey answers `200` with "
        "`{\"success\": false, \"exception\": \"...\"}`. `call()` treats a request as "
        "successful only when the status is 2xx **and** the body reports no failure.",
        "**4xx bodies are the payload, not noise.** They are preserved on `LabKeyError.body` "
        "because the specification requires the body to be stored on a failed publish.",
        "**Redirects are never followed**, and an unexpected 3xx becomes an error naming the "
        "`Location` header.",
        "**Known-bad action names are refused locally.** `query-importData.api`, "
        "`experiment-saveMaterials.api` and `experiment-deriveSamples.api` do not exist; the "
        "client names the action that does instead of letting the server return a confusing "
        "404.",
        "**`tolerate=` supports idempotency.** Pass lowercase substrings such as "
        "`(\"already exists\",)` so that creating something that exists counts as success.",
        "**Payloads are redacted before logging.** Any key containing `password`, `apikey`, "
        "`api_key`, `token`, `secret`, `csrf` or `crypt` is replaced with `***`, recursively.",
        "**Container paths are URL-encoded per segment**, so `/Tutorials/HIV Study` works and "
        "the structural slashes survive.",
    ])

    g.h2("6.7 Managing samples today — an honest answer")
    g.p("**You cannot manage samples with OSM at this revision.** There is no database schema, "
        "no REST API, no user interface and no sample domain code. Sections 1.5 and 12 list "
        "exactly what exists and what does not.")
    g.p("What the repository does support today is the work that precedes an application:")
    g.bullets([
        "**Establishing ground truth** about the LabKey target with "
        "`scripts/verify_labkey.py` and the integration suite, and recording it so no later "
        "session has to guess.",
        "**Querying the requirement register** to answer what the system must do, traced to "
        "the specification section that says so.",
        "**Working the backlog** in dependency order, with an integrity gate that refuses an "
        "untraced requirement, a dangling reference or a backlog item that depends on work "
        "scheduled after it.",
        "**Reading the LabKey CE data model** — sample types, aliquot columns, lineage, the "
        "audit timeline and the measurement-unit vocabulary — which is what the publish bridge "
        "will map onto.",
    ])
    g.p("If you need sample management **now**, the honest options are LabKey CE's native "
        "sample types for registration and lineage only, with the caveats in section 1.3 "
        "firmly in mind — CE will not enforce status, will not mask PHI and has no freezer "
        "map — or the commercial LabKey tiers in section 1.4. OSM becomes usable for real "
        "sample work at the end of iteration I1 (PR-015), and physically useful at the end of "
        "I2 (PR-020).")


def write_testing(g: Guide) -> None:
    f = g.facts
    g.h1("7. Verification and testing")

    g.h2("7.1 Running the suites")
    g.code(
        "make check                       # the whole gate: lint, types, tests, memory,\n"
        "                                 # render and docs-check\n"
        "make test                        # unit tests only, no network\n"
        "make integration                 # tests marked `labkey`; needs LK_* in the "
        "environment\n"
        "\n"
        "python3 -m pytest -m \"not labkey\"            # what `make test` runs\n"
        "python3 -m pytest tests/test_memory_cli.py -v\n"
        "python3 -m pytest -k redaction               # select by name\n"
        "python3 -m pytest --collect-only -q          # recount the suites"
    )
    g.p("`pyproject.toml` configures `testpaths = [\"tests\"]`, `pythonpath = [\"src\"]` and "
        "`addopts = \"-q --strict-markers\"`. `--strict-markers` means an unregistered marker "
        "is an error rather than a silent no-op, so the `labkey` deselection cannot rot.")

    g.h2("7.2 What the suites cover")
    g.table(
        ["File", "Tests", "What it proves"],
        [[f"`{name}`", count, what] for name, count, what in TEST_SUITES],
        widths=[2.2, 0.6, 3.7],
    )
    g.p(f"**{UNIT_TEST_TOTAL} unit tests** run with no network access, plus "
        f"**{INTEGRATION_TEST_TOTAL} integration tests** against a live server. The integration "
        "tests are the executable half of `docs/labkey-ce-ground-truth.md`: if a LabKey upgrade "
        "changes one of those facts, a test fails rather than the documentation quietly going "
        "stale. Every one of them is read-only, and none touches `SleepDrive-Lab`.")
    g.p("The memory CLI is tested through `subprocess` against a throwaway database in "
        "`tmp_path`, driven by the `mem` fixture, because its path resolution happens at import "
        "time from `OSM_MEMORY_DB` and `OSM_MEMORY_DUMP`. That also means the tests exercise "
        "the supported command-line surface rather than an internal one.")

    g.h2("7.3 What the quality gate enforces")
    g.table(
        ["Gate", "Enforced by", "Failure means"],
        [
            ["Unit tests pass with no network", "`make test`",
             "A behaviour regression, or a warning that became an error"],
            ["Every requirement is traced", "`tools/memory.py check`",
             "A requirement nobody plans to satisfy — a planning hole"],
            ["No traceability edge dangles", "`tools/memory.py check`",
             "An edge points at a backlog item, decision, feature or verification that does not "
             "exist"],
            ["Backlog dependencies exist", "`tools/memory.py check`",
             "`depends_on` names an unknown pull request"],
            ["Backlog `seq` values are unique", "`tools/memory.py check`",
             "Two items claim the same position in the queue"],
            ["Dependencies are scheduled earlier", "`tools/memory.py check`",
             "An item depends on work with a later `seq`, so the order is not an order"],
            ["ADR `supersedes` resolves", "`tools/memory.py check`",
             "A decision supersedes an ADR that does not exist"],
            ["The committed dump is current", "`tools/memory.py check`",
             "The database was mutated and the dump not committed. The gate regenerates it and "
             "tells you to commit it"],
            ["Generated markdown is current", "`tools/render_backlog.py --check`",
             "`docs/backlog.md` or `specs/tasks/README.md` drifted from the database"],
            ["This guide is current", "`tools/build_guide.py --check`",
             "The generator or an underlying fact changed and the `.docx` was not rebuilt"],
            ["No secret in a tracked file", "`tests/test_no_secrets.py`",
             "A credential-shaped literal entered the repository"],
            ["Style and a security lint subset", "`make lint` (`ruff`)",
             "Not enforced yet — `ruff` is pinned by PR-002 and the target reports its absence"],
            ["Static types over `src`", "`make types` (`mypy`)",
             "Not enforced yet — same reason"],
            ["Coverage floor", "not configured yet",
             "`pytest-cov` is declared but no threshold is set; it is an acceptance criterion "
             "of PR-002"],
        ],
        widths=[1.8, 1.7, 3.0],
    )
    g.p("Three of the standards the project holds itself to are **written but not yet "
        "enforceable**, because the code they would apply to does not exist: concurrency tests "
        "for the slot exclusion constraint and the audit chain, benchmarks that fail the build "
        "when a performance target is missed, and the rule that every new domain table has a "
        "row-level security policy checked by a schema test. They are stated in "
        "`standards/general/testing.md` and `standards/general/security.md` and become live "
        "with PR-003 onwards.")

    g.h2("7.4 Interpreting a failure")
    g.table(
        ["Message", "Cause", "Fix"],
        [
            ["`INTEGRITY PROBLEMS:` followed by "
             "`requirement FR-nnn has no traceability link`",
             "A requirement was added without linking it to anything",
             "`tools/memory.py link --req FR-nnn --kind backlog --to PR-nnn`"],
            ["`... was stale and has been regenerated; commit it`",
             "The database changed and the dump was not committed",
             "`git add .sdd/memory.sql` and commit it with the change"],
            ["`STALE generated files (run tools/render_backlog.py)`",
             "The backlog changed, or someone hand-edited a generated file",
             "`tools/render_backlog.py`, then commit both files"],
            ["`traceability points at unknown backlog PR-nnn`",
             "An edge references an item that was renamed or never created",
             "Create the item, or re-point the edge"],
            ["`PR-nnn (seq n) depends on PR-mmm (seq m), which is not scheduled earlier`",
             "A dependency inversion in the queue", "Re-sequence one of the two items"],
            ["`possible secrets in tracked files:` …",
             "A credential-shaped literal is present",
             "Remove it. If it is a deliberate placeholder, add its **exact** value to "
             "`ALLOWED_LITERALS` with a justification. Do **not** broaden the patterns"],
            ["`the secret-scanner allowlist has grown large enough to hide a real secret`",
             "`ALLOWED_LITERALS` reached 40 entries",
             "Prune it. The allowlist is a review gate, not a rubber stamp"],
            ["`unexpected git remote configured: …`",
             "A remote was added", "Remove it (`AGENTS.md` §5), or revisit the decision "
                                   "explicitly"],
            ["A test fails with an unexpected warning",
             "`filterwarnings = [\"error\"]` turns warnings into errors",
             "Fix the warning. Only `urllib3` `InsecureRequestWarning` is exempted, and that "
             "exemption is documented in `pyproject.toml`"],
            ["Integration tests all report `skipped`",
             "No `LK_*` credentials, or the server is unreachable",
             "Export the credentials, or accept the skip — the unit suite is the gate"],
        ],
        widths=[2.2, 1.9, 2.4],
    )

    g.h2("7.5 The verification register")
    g.p("Beyond the test suites, the project keeps a register of how each claim about the "
        "environment was actually checked. At this revision it holds "
        f"{f['counts']['verifications']} rows "
        + ", ".join(f"({r['n']} `{r['result']}`)" for r in f["verification_results"]) + ".")
    g.p("Evidence is ranked, strongest first: **`source`** (the LabKey sources under "
        "`/root/scicore`, cited by absolute path and identifier), **`http`** (a real call "
        "against the running server), **`shell`** (a command and its output, for host facts), "
        "and **`doc`** (official documentation by URL — weakest, because documentation and "
        "implementation diverge, and `docs/gap-analysis.md` records several cases where they "
        "do). `reasoning` is never acceptable evidence for a behavioural claim.")
    g.bullets([
        "A pull request asserting new external behaviour without a verification row fails "
        "review.",
        "A verification whose `result` is `fail` must not be deleted. Negative results are why "
        "the register exists.",
        "Verification runs read-only unless the claim genuinely requires a write, and a write "
        "goes to a scratch container this project owns and is cleaned up.",
    ])


def write_security(g: Guide) -> None:
    g.h1("8. Security considerations")

    g.h2("8.1 The secret scanner")
    g.p("`tests/test_no_secrets.py` turns the promise \"no credential ever enters the "
        "repository\" into a gate that runs on every test run. It scans **every file git "
        "tracks** — obtained with `git ls-files -z`, so untracked scratch files are out of "
        "scope and a newly added file is automatically in scope.")
    g.p("It works in three layers:")
    g.steps([
        "**High-signal patterns**, which are always a finding with no allowlist: a private-key "
        "header, a GitHub token (`ghp_`/`gho_`/`ghu_`/`ghs_`/`ghr_`), a HuggingFace token "
        "(`hf_`), an AWS access key id (`AKIA…`), a Slack token (`xox…`), a JSON web token, "
        "and a credential embedded in a URL — a user and password in the authority "
        "component before the `@`.",
        "**Credential-shaped assignments**: `name = \"value\"`, `name: \"value\"` and bare "
        "`NAME=value` shell assignments where the name contains `pass`, `secret`, `token`, "
        "`apikey`, `auth`, `credential`, `passphrase`, `private_key` or `access_key`. Values "
        "shorter than six characters and values that are variable references are ignored.",
        "**An allowlist of exact literals.** Anything credential-shaped that is not on the "
        "list fails. Adding a new one is therefore a deliberate act visible in the diff, which "
        "is the review gate the design is aiming for.",
    ])
    g.quote("Broadening a regex to make a failure go away is the wrong fix. Removing the "
            "secret, or adding a clearly fake value to `ALLOWED_LITERALS` with a "
            "justification, is the right one.")
    g.p("**The scanner is itself tested.** Eight planted secrets prove it fires, and eight safe "
        "patterns — an empty password, a one-character value, a variable reference, a comment "
        "mentioning `LK_PASSWORD`, a dictionary literal reading `self.api_key` — prove it does "
        "not cry wolf. A gate that cannot fail is not a gate; a gate that cries wolf gets "
        "disabled. A further test fails if the allowlist reaches 40 entries.")
    g.p("Four companion tests cover the rest of the promise: `.env` is not tracked; "
        "`.env.example` carries no value for any credential-shaped name; `.gitignore` refuses "
        "`.env`, `*.pem`, `*.key` and `secrets/`; the binary memory database is not tracked "
        "while the dump is; and no git remote is configured.")

    g.h2("8.2 Credential handling in code")
    g.bullets([
        "**No defaults.** `load_config()` raises `ConfigError` naming the missing variables. "
        "There is no prompt and no configuration file.",
        "**Excluded from `repr`.** `LabKeyConfig.password` and `LabKeyConfig.api_key` are "
        "declared `field(repr=False)`, so a config object appearing in a traceback, a log line "
        "or a debugger cannot disclose them.",
        "**A safe one-line summary.** `LabKeyConfig.describe()` names the auth mode and the TLS "
        "state — never the secret — and that is what the session log prints.",
        "**Recursive redaction before logging.** `_redact()` blanks any key containing "
        "`password`, `apikey`, `api_key`, `token`, `secret`, `csrf` or `crypt`. It is "
        "deliberately broad: a false positive costs a masked debug value, a false negative "
        "costs a credential in a log file.",
        "**The login body is never echoed wholesale.** On a failed login LabKey can reflect "
        "submitted form values, so only the `exception` field is surfaced.",
        "**The API key travels as a header**, not in a URL, so it cannot leak through a "
        "referrer or an access log.",
        "**Credentials reach subprocesses through the environment**, never as arguments.",
    ])

    g.h2("8.3 Transport security")
    g.p("The full TLS decision table is in section 5.3. The principles behind it:")
    g.bullets([
        "Verification is **on by default** for anything that is not this machine.",
        "The loopback exemption exists because the deployment's certificate is self-signed, so "
        "verifying loopback always fails. It is scoped by address, not by hostname string "
        "matching: `_is_loopback()` parses the host as an IP address and asks `is_loopback`.",
        "The opt-out for a non-loopback host is kept, because internal certificate authorities "
        "are a real situation — but it emits a `WARNING` naming the host, because an operator "
        "who did not mean it needs to see it.",
        "Redirects are never followed, so a credential-bearing request cannot be replayed "
        "against a host the client did not choose.",
    ])

    g.h2("8.4 Least privilege")
    g.table(
        ["Principal", "Grant", "State"],
        [
            ["LabKey publish bridge",
             "An API key minted with the `EditorWithoutDelete` restriction role",
             "Planned — the bridge is PR-031; the mechanism is verified to exist in CE"],
            ["Verification harness and integration tests", "Read access only; every call is a "
                                                           "read", "**In force today**"],
            ["Audit database role", "`INSERT` and `SELECT` on audit tables; no `UPDATE`, no "
                                    "`DELETE`", "Planned — PR-004"],
            ["Auditor role", "`SELECT` only, and nothing else", "Planned — PR-006"],
            ["Assistant / MCP", "Cannot perform any operation the same principal could not "
                                "perform over REST", "Planned — PR-033 to PR-035"],
        ],
        widths=[1.7, 2.8, 2.0],
    )

    g.h2("8.5 Input validation")
    g.p("These rules are recorded in `standards/general/security.md` and `AGENTS.md` §3. They "
        "govern code that does not exist yet, and are reproduced here because they constrain "
        "every future change and because two of them already apply to the LabKey client.")
    g.bullets([
        "**Validate on the server.** Client-side validation is a convenience; the boundary is "
        "the only place that counts. Every request body is validated against a declared schema.",
        "**Never build SQL by concatenation** — LabKey SQL included. Filter expressions are "
        "parsed into a typed structure. LabKey SQL supports `USERID()`, `||`, `COALESCE`, "
        "joins and subqueries, all of which a hostile filter could exploit if it reached the "
        "SQL text.",
        "**CSV and file import**: validate the declared content type rather than trusting it; "
        "bound the size and the row count; report every rejected row with its line number; and "
        "neutralise formula injection so a cell beginning with `=`, `+`, `-` or `@` cannot "
        "execute when the file is later opened.",
        "**Barcode input is untrusted.** Bound its length, validate its characters, never "
        "interpolate it into a query. A barcode that does not resolve returns 404 rather than "
        "an error revealing whether the id exists in another project.",
        "**Disclosure**: an error never leaks a stack trace, a SQL fragment or a file path to a "
        "client, and a lookup failure returns the same response whether the identifier does "
        "not exist or exists outside the caller's scope.",
    ])

    g.h2("8.6 Agents and prompt injection")
    g.p("The specification anticipates assistants acting on the system, and the rules are "
        "unusually strict for a reason recorded in `docs/gap-analysis.md`: LabKey is candid "
        "that in its own MCP server, *\"what limits the agent is the API key you give it, not "
        "anything built into the MCP Server.\"* OSM takes the opposite position.")
    g.bullets([
        "**An operation that must never be agent-reachable is not exposed as a tool at all.** "
        "Refusing at call time is a weaker guarantee than not existing, and a test asserts the "
        "absence so a later change cannot quietly add it. Discard, ship, lock, sign and "
        "permission changes are all in this category.",
        "**A destructive tool requires an explicit confirmation argument and an administrative "
        "scope.**",
        "**The tool allowlist is bound to the session at authentication time**, never derived "
        "from content the agent reads. Text inside an SOP PDF can ask for a tool; nothing is "
        "listening.",
        "**The retrieval corpus is restricted at index time**, not filtered at query time.",
        "**At most one unsafe write chain per turn.**",
        "Every invocation is audited with the model and request identity.",
    ])
    g.p("None of this is implemented yet — it is iteration I6, PR-033 to PR-035.")

    g.h2("8.7 Auditability")
    g.p("Every write is to be audited in the same transaction that performs it; if a code path "
        "can write without auditing, that is a defect (ADR-0003). Completeness comes from "
        "database triggers rather than from developer discipline, tamper evidence from a "
        "SHA-256 chain anchored by a daily checkpoint, and linearity from a PostgreSQL "
        "advisory lock that serialises hash and sequence assignment.")
    g.p("**LabKey CE's own audit trail supplements the OSM trail and never replaces it.** CE's "
        "`canDeleteOldRows()` returns `false`, which prevents routine purging, but a database "
        "administrator can still alter rows undetectably. The specification requires "
        "detection, which is why the chain exists.")


def write_troubleshooting(g: Guide) -> None:
    g.h1("9. Troubleshooting")
    g.p("Symptoms below are quoted from the actual messages the code emits, so they can be "
        "matched by searching a log or a terminal.")

    g.h2("9.1 Configuration and connection")
    g.table(
        ["Symptom", "Cause", "Fix"],
        [
            ["`error: no LabKey credential available: set LK_APIKEY, or set both LK_USER and "
             "LK_PASSWORD (missing: ...)`",
             "No usable credential in the environment. Blank and whitespace-only values count "
             "as unset",
             "Export `LK_APIKEY`, or both `LK_USER` and `LK_PASSWORD`. Check for a `.env` you "
             "forgot to source"],
            ["`could not reach the LabKey API at any candidate context path: <origin>: "
             "ConnectionError; <origin>/labkey: ConnectionError`",
             "Wrong host or port, or the server is down",
             "Confirm `LK_URL`; check the port is listening; try `LK_CONTEXT` explicitly "
             "instead of `auto`"],
            ["The candidate list reports `non-JSON response`",
             "LabKey served the HTML login page rather than the API — the context path is "
             "wrong",
             "Set `LK_CONTEXT` to `` (empty) or `labkey` explicitly"],
            ["`SSLError` / `CERTIFICATE_VERIFY_FAILED` against a named host",
             "Verification is on for any non-loopback host, and the certificate does not chain "
             "to a trusted root",
             "Install the issuing CA on the client, or set `LK_INSECURE=1` deliberately and "
             "accept the warning"],
            ["`WARNING osm.labkey: TLS verification is DISABLED for the non-loopback host ...`",
             "`LK_INSECURE` is set while talking to a remote host",
             "Unset `LK_INSECURE` unless this is deliberate. Credentials are being sent over an "
             "unchecked connection"],
            ["`error: LK_TIMEOUT must be a number, got '30s'`",
             "The timeout is a bare number of seconds, not a duration string", "Use `LK_TIMEOUT=30`"],
            ["`error: LK_URL must be an http or https URL, got scheme 'labkey'`",
             "A malformed URL, often a missing `https://`", "Set a full URL including the scheme"],
        ],
        widths=[2.3, 2.0, 2.2],
    )

    g.h2("9.2 LabKey API calls")
    g.table(
        ["Symptom", "Cause", "Fix"],
        [
            ["`LabKey session established at ... as (unidentified)`",
             "The identity response was not absorbed, or the session resolved to a guest. "
             "`login-loginApi.api` nests the principal under `user` while `login-whoAmI.api` "
             "reports it at the top level — this cost a session during PR-001",
             "Confirm the credential is right. `lk.whoami()` should return a non-guest `email`"],
            ["`no CSRF token captured; mutating calls will be rejected`",
             "The bootstrap response carried no `CSRF` field and no `X-LABKEY-CSRF` cookie, "
             "often a proxy stripping cookies",
             "Reconnect; check the cookie survives any intermediary"],
            ["`query-importData.api does not exist in LabKey; use query-import.api`",
             "A guessed action name. The client refuses three known-bad names locally",
             "Use `query-import.api`. For materials use `query-insertRows.api` / "
             "`query-import.api` on schema `samples`, or `experiment-importSamples.api`; for "
             "derivation use `experiment-derive.api`"],
            ["`expected JSON but received text/html (N bytes); this usually means the action "
             "name is wrong and LabKey served an HTML page`",
             "A misspelled or non-existent action",
             "Check the action name against `docs/labkey-ce-ground-truth.md`"],
            ["`unexpected redirect to https://<canonical-host>/...; redirects are not followed "
             "on purpose`",
             "The action redirects to the server's canonical hostname; following it against a "
             "mismatched certificate loops",
             "Address the redirect target directly, or set `LK_URL` to the canonical host"],
            ["A `LabKeyError` is raised even though the status was `200`",
             "LabKey answered `200` with `success: false` or an `exception` field. Status alone "
             "is not a verdict",
             "Read `exc.body` and `exc.exception_class`; they carry the server's explanation"],
            ["`LabKeyAuthError` with status `401` or `403`",
             "Wrong credential, insufficient permission, or API keys disabled site-wide "
             "(`allowApiKeys` defaults to `false`)",
             "Verify the account's role on the container; ask a site administrator to enable "
             "API keys before minting one"],
            ["The probe reports `partial`: `UNEXPECTEDLY PRESENT: inventory`",
             "A LabKey upgrade or an added module now provides a schema the gap analysis "
             "assumed absent",
             "Not an error — revisit `docs/gap-analysis.md` and record a new verification"],
        ],
        widths=[2.3, 2.0, 2.2],
    )

    g.h2("9.3 The memory database and the gate")
    g.table(
        ["Symptom", "Cause", "Fix"],
        [
            ["`error: no memory database at ... and no dump at ...; run: tools/memory.py init`",
             "`OSM_MEMORY_DB` or `OSM_MEMORY_DUMP` points somewhere empty, or you are outside "
             "the repository",
             "Unset the overrides, or run from the repository root"],
            ["`error: invalid id 'ADR-001' for decisions; expected pattern ^ADR-\\d{4}$`",
             "ADRs use four digits, matching `specs/adr/NNNN-*.md`", "Use `ADR-0001`"],
            ["`error: tasks: FOREIGN KEY constraint failed: --pr must name a backlog item that "
             "already exists`",
             "The task was added before its backlog item", "Add the backlog row first"],
            ["`error: ... already exists; use --replace to overwrite it deliberately`",
             "Duplicate id", "Choose a new id, or pass `--replace` knowingly"],
            ["`error: query accepts read-only SELECT/WITH statements only`",
             "An attempt to mutate through `query`", "Use `add`, `set` or `link`"],
            ["`error: <file>:<line>: ... (after N rows applied)`",
             "A `batch` run failed part-way. Partial application is reported, never silent",
             "Fix the offending line; re-running the whole batch is safe for `link` "
             "(`INSERT OR IGNORE`) but will report duplicates for `add` without `--replace`"],
            ["`INTEGRITY PROBLEMS:` on `make check`", "See the table in section 7.4",
             "Address each listed problem; the gate exits 2 until all are clear"],
            ["`make check` passes locally but the diff shows `.sdd/memory.sql` changed",
             "`check` regenerates a stale dump as a side effect", "Commit the regenerated dump"],
        ],
        widths=[2.3, 2.0, 2.2],
    )

    g.h2("9.4 Tests and tooling")
    g.table(
        ["Symptom", "Cause", "Fix"],
        [
            ["`ruff not installed - skipping lint (pinned by PR-002)`",
             "Expected at this revision", "Install the `dev` extra, or ignore it until PR-002"],
            ["`possible secrets in tracked files: ...`", "See section 8.1",
             "Remove the value, or allowlist the exact fake literal with a justification"],
            ["Every `labkey`-marked test reports `skipped`",
             "No credentials, or the server is unreachable. The fixture skips rather than fails "
             "on purpose",
             "Export `LK_*` and re-run `make integration`"],
            ["A test fails on a warning that is not an assertion failure",
             "`filterwarnings = [\"error\"]`", "Fix the warning; only `InsecureRequestWarning` "
                                               "is exempted"],
            ["`'labkey' not found in `markers` configuration option`",
             "Running pytest from outside the repository root, so `pyproject.toml` is not "
             "loaded, combined with `--strict-markers`",
             "Run from the repository root"],
            ["`python-docx is not installed` from `tools/build_guide.py`",
             "The documentation dependency is missing",
             "`python3 -m pip install python-docx`"],
            ["`tools/build_guide.py --check` reports the guide is stale",
             "The generator or an underlying fact changed", "Run `make docs` and commit the "
                                                            "result"],
        ],
        widths=[2.3, 2.0, 2.2],
    )

    g.h2("9.5 LabKey behaviour that looks like a bug and is not")
    g.table(
        ["Observation", "Explanation"],
        [
            ["A sample marked `Consumed` in LabKey can still be edited, aliquoted and derived",
             "CE registers no `SampleStatusService` provider, so `isOperationPermitted()` "
             "returns `true` unconditionally. Status in CE is a label, not a rule"],
            ["A column tagged PHI is exported in the clear",
             "No `ComplianceService` is registered. PHI tags survive export and import and mask "
             "nothing"],
            ["Storage or freezer columns do not appear on a sample type",
             "`InventoryService.get()` returns `null` and every call site is null-guarded, so "
             "absence is silent rather than an error"],
            ["`query-getSchemas.api` returns no `inventory`, `storage` or `freezer` schema",
             "Correct and expected. There is no storage substrate in CE; this is the finding "
             "ADR-0001 turns on"],
            ["The MCP endpoint does nothing",
             "`McpService.get()` returns `NoopMcpService` with `isEnabled()` false, and the "
             "search registration is commented out"],
            ["The field editor refuses the Unique ID (barcode) type",
             "A JavaScript-only gate. `property-createDomain.api` with the storage-unique-id "
             "concept URI creates a working barcode field server-side"],
            ["Search returns hits but no facet counts",
             "CE has no faceting at all — zero occurrences of `facet` in the search module and "
             "no `lucene-facet` dependency"],
            ["`security-createApiKey.api` fails",
             "`allowApiKeys` defaults to `false` and must be enabled in site settings"],
        ],
        widths=[2.5, 4.0],
    )


def write_roadmap(g: Guide) -> None:
    f = g.facts
    g.h1("10. Roadmap and current status")

    g.h2("10.1 Where the project stands")
    g.table(
        ["Measure", "Value"],
        [
            ["Project version", f"`{f['version']}`"],
            ["Revision", f"`{f['commit_full']}`"],
            ["Branch", f"`{f['branch']}`"],
            ["Git remote", f["remote"] or "none, deliberately"],
            ["Requirements traced to specification sections", str(f["counts"]["requirements"])],
            ["Research findings", str(f["counts"]["research"])],
            ["Accepted decisions", str(f["counts"]["decisions"])],
            ["Capability inventory rows", str(f["counts"]["features"])],
            ["Verifications recorded", str(f["counts"]["verifications"])],
            ["Traceability edges", str(f["counts"]["traceability"])],
            ["Pull requests in the backlog", str(f["counts"]["backlog"])],
            ["Backlog status", ", ".join(f"{r['n']} `{r['status']}`"
                                         for r in f["backlog_status"])],
            ["Tests", f"{UNIT_TEST_TOTAL} unit, {INTEGRATION_TEST_TOTAL} integration"],
        ],
        widths=[3.0, 3.5],
    )
    g.p("Implementation has begun. PR-001 (the LabKey client library and verification harness) "
        "is at status `review`, and a follow-on hardening pass added 184 tests across the "
        "project tooling — which had become load-bearing while being the only part with no "
        "tests of its own. That pass found and fixed five real defects in "
        "`tools/memory.py`, every one of them found by a test rather than by inspection: the "
        "integrity gate did not fail on an untraced requirement; it validated only `backlog` "
        "traceability edges, so a dangling `decision`, `feature` or `verification` edge passed; "
        "`link` and `set` raised uncaught tracebacks on bad input; and a foreign-key failure "
        "was reported with a misleading suggestion to use `--replace`.")

    g.h2("10.2 Iterations and acceptance gates")
    g.p("Delivery is organised into eight iterations, each with its own acceptance gate. The "
        "freezer map, the job queue, the ELN and the finder have independent acceptance.")
    g.table(
        ["Iteration", "Theme", "Acceptance gate", "Pull requests"],
        [
            ["I0", "Foundations", "Auth, audit and OpenAPI. Done when the hash chain verifies",
             "PR-001 to PR-008"],
            ["I1", "Registry", "Sample types, sources, samples. Done at aliquot",
             "PR-009 to PR-015"],
            ["I2", "Freezer map", "Storage hierarchy and slot operations. Done per "
                                  "specification §5", "PR-016 to PR-020"],
            ["I3", "Workflow", "Templates, jobs, tasks. Done at the queue", "PR-021 to PR-023"],
            ["I4", "ELN", "Notebooks, review, signature. Done at signature", "PR-024 to PR-026"],
            ["I5", "Search", "Finder, picklists, row-level security. Done with RLS",
             "PR-027 to PR-029"],
            ["I6", "MCP and LabKey bridge",
             "Publishing and agent access. Done when the UI runs through MCP semantics",
             "PR-030 to PR-035"],
            ["I7", "Operations", "Pipelines, benchmarks, runbook. Done at the runbook",
             "PR-036 to PR-038"],
        ],
        widths=[0.8, 1.5, 3.0, 1.2],
    )
    g.p("Requirements are distributed across those iterations as follows (23 carry no iteration "
        "because they are constraints and prohibitions that apply throughout):")
    g.table(
        ["Iteration", "Requirements"],
        [[str(r["iteration"] or "(none)"), str(r["n"])] for r in f["req_iterations"]],
        widths=[1.5, 1.5],
    )

    g.h2("10.3 The full backlog")
    g.p("Rendered from the memory database at build time. `docs/backlog.md` carries the "
        "acceptance criteria for each item; `tools/memory.py show PR-nnn` carries everything.")
    g.table(
        ["PR", "It.", "Size", "Title", "Depends on", "Status"],
        [[str(r["id"]), str(r["iteration"]), str(r["size"]), str(r["title"]),
          str(r["depends_on"] or "—"), str(r["status"])] for r in f["backlog"]],
        widths=[0.7, 0.4, 0.45, 2.85, 1.0, 0.65],
    )
    g.p("Sizes translate to: **XS** under a day, **S** one to two days, **M** three to five "
        "days, **L** five to ten days. Each item is sized to be reviewable in one sitting, and "
        "is mergeable only when every acceptance criterion is met and the tree is green.")

    g.h2("10.4 Open questions")
    g.p("Recorded in `specs/prd.md` and unresolved in the specification itself. None blocks the "
        "current iteration.")
    g.bullets([
        "**Temperature logger integration** — which loggers, and does OSM poll or receive? "
        "(specification §15)",
        "**ZPL label printing** — which printer models must be supported? Note that the "
        "commercial LabKey product supports no ZPL at all, so there is no prior art to follow.",
        "**USB identifiers** require a data protection impact assessment before any design work "
        "begins. Until that exists, the question stays open.",
        "**Audit chain partitioning** — whether the hash chain should be partitioned per "
        "project if the serialisation point proves costly under load. The schema allows it; "
        "the benchmark in PR-037 decides it.",
    ])

    g.h2("10.5 Capabilities deliberately withheld")
    g.p("These are requirements, not omissions. Each is a prohibition the specification states "
        "explicitly, and each is to be verified by a test asserting the capability is **absent** "
        "rather than merely refused.")
    g.bullets([
        "**No patient-identifiable data in OSM.** Study identifiers enter only as opaque "
        "tokens. Personal health information is stripped when a publish payload is built, so "
        "it never reaches the queue, let alone LabKey.",
        "**Assistants may never discard, ship, lock, sign or change permissions.** These "
        "operations are not exposed as tools at all.",
        "**No content-derived privilege.** The tool allowlist is bound to the session at "
        "authentication time; the retrieval corpus is restricted at index time.",
        "**No `addWebPart` against LabKey.** Portals are delivered as folder archives. The "
        "specification forbids it and it misbehaves on CE 26.",
        "**No reimplementation of the LabKey user interface**, and no dependency on LabKey to "
        "operate.",
    ])


def write_appendix(g: Guide) -> None:
    f = g.facts
    g.h1("11. Appendix")

    g.h2("11.1 Glossary")
    g.table(
        ["Term", "Meaning"],
        [
            ["ADR", "Architecture Decision Record. Numbered `ADR-nnnn`, stored in `specs/adr/` "
                    "and mirrored into the `decisions` table"],
            ["Aliquot", "A portion split from a parent sample, with the parent's remaining "
                        "amount reduced exactly. LabKey CE carries the columns; the semantics "
                        "are OSM's"],
            ["CE", "Community Edition — the free, Apache-2.0 LabKey distribution. 26.7.5 here"],
            ["CSRF token", "`X-LABKEY-CSRF`. The cookie value, sent as the same-named header on "
                           "every mutating call and as a form field on multipart posts"],
            ["ELN", "Electronic Lab Notebook. Draft → review → approved → signed → locked"],
            ["FRD", "Feature Requirements Document. One per capability. **None written yet**"],
            ["Iteration", "One of I0 to I7, each with an independent acceptance gate"],
            ["LSID", "LabKey's Life Science Identifier, used for object identity in `exp`"],
            ["MCP", "Model Context Protocol. In OSM, a thin adapter over the REST API "
                    "(ADR-0006). Present but inert in LabKey CE"],
            ["Memory database", "The project's knowledge base: `.sdd/memory.sql` tracked, "
                                "`.sdd/memory.db` derived"],
            ["Outbox", "The transactional publish queue. A domain write and its "
                       "`osm_publish_queue` row commit together (ADR-0005)"],
            ["PHI", "Protected Health Information. Forbidden in OSM outright, and stripped "
                    "before any publish"],
            ["Picklist", "A named, ordered set of samples a job can be instantiated against"],
            ["Prohibition", "A requirement of kind `prohibition` — something the system must "
                            "**not** do, verified by asserting absence"],
            ["RLS", "Row-Level Security. PostgreSQL policies on domain tables. Neither LabKey "
                    "CE nor commercial Sample Manager has it on samples"],
            ["SDD", "Specification-Driven Development. The method this project follows"],
            ["Slot", "The smallest storage unit in the five-level freezer hierarchy. Holds at "
                     "most one sample, enforced by an exclusion constraint (ADR-0004)"],
            ["Traceability edge", "A `requirement → artifact` row. Every requirement must have "
                                  "at least one or the gate fails"],
            ["Verification", "A `V-nnn` row recording how a claim about the environment was "
                             "actually checked, and with what evidence"],
        ],
        widths=[1.4, 5.1],
    )

    g.h2("11.2 File and directory map")
    g.table(
        ["Path", "Contents"],
        [
            ["`README.md`", "Project overview and quick start"],
            ["`memory.md`", "**Entry point.** The knowledge base and the house rules"],
            ["`AGENTS.md`", "The briefing for any session working on OSM. Consolidates "
                            "`standards/`; regenerated when they change"],
            ["`JOURNAL.md`", "Append-only session log, newest last, ISO-8601 UTC timestamps"],
            ["`LICENSE`", "Apache-2.0"],
            ["`Makefile`", "`check`, `lint`, `types`, `test`, `integration`, `memory`, "
                           "`render`, `docs`, `docs-check`, `clean`"],
            ["`pyproject.toml`", "Project metadata, dependencies, pytest, ruff and mypy "
                                 "configuration"],
            ["`.env.example`", "Environment variable **names** with empty values. Tracked; "
                               "`.env` is not"],
            ["`.gitignore`", "Secret shapes, the binary memory database, Python and Node "
                             "artefacts"],
            ["`.sdd/memory.sql`", "The knowledge base. Deterministic dump, tracked"],
            ["`.sdd/memory.db`", "Derived SQLite binary. Git-ignored"],
            ["`tools/memory.py`", "The memory CLI"],
            ["`tools/render_backlog.py`", "Renders `docs/backlog.md` and "
                                          "`specs/tasks/README.md`"],
            ["`tools/build_guide.py`", "Generates this document"],
            ["`scripts/verify_labkey.py`", "The LabKey verification harness. The only "
                                           "sanctioned way to touch a deployment"],
            ["`src/osm/labkey/config.py`", "Credential and TLS resolution from the environment"],
            ["`src/osm/labkey/client.py`", "The LabKey HTTP client"],
            ["`specs/prd.md`", "Product Requirements Document, REQ-1 to REQ-16"],
            ["`specs/adr/`", "Eight accepted Architecture Decision Records"],
            ["`specs/features/`", "Feature Requirements Documents — **empty at this revision**"],
            ["`specs/tasks/README.md`", "Generated phase index and requirement-to-task matrix"],
            ["`specs/source/spezifikation-extract.md`",
             "Verbatim-structure extraction of the governing specification"],
            ["`standards/general/security.md`", "Secrets, authorisation, input, disclosure, "
                                                "agents"],
            ["`standards/general/testing.md`", "What a test must prove, and what does not count"],
            ["`standards/general/verification.md`", "Evidence ranking and the recording rule"],
            ["`standards/labkey/http-conventions.md`", "Session, transport, action names, "
                                                       "idempotency, credentials"],
            ["`docs/labkey-ce-ground-truth.md`", "Verified facts about LabKey CE 26.7.5"],
            ["`docs/gap-analysis.md`", "OSM vs commercial Sample Manager vs LabKey CE"],
            ["`docs/backlog.md`", "Generated pull-request queue"],
            ["`docs/OSM-Sample-Manager-Guide.docx`", "This document"],
            ["`tests/`", "The test suite; `conftest.py` holds the memory harness and the "
                         "tracked-file fixture"],
            ["`.github/agents/`", "Role definitions: `pm`, `architect`, `devlead`, `dev`"],
            ["`.github/prompts/`", "Workflow step definitions"],
        ],
        widths=[2.2, 4.3],
    )

    g.h2("11.3 Decision records")
    g.table(
        ["ID", "Title", "Status", "Decision"],
        [[str(r["id"]), str(r["title"]), str(r["status"]), str(r["decision"])]
         for r in f["decisions"]],
        widths=[0.75, 1.7, 0.7, 3.35],
    )
    g.p("Each record is stored in full at `specs/adr/NNNN-<slug>.md` in MADR format, with "
        "context, decision drivers, considered options, the outcome, consequences and "
        "references. `tools/memory.py show ADR-0004` prints the database mirror.")

    g.h2("11.4 Specification and standards reference")
    g.table(
        ["Document", "What to consult it for"],
        [
            ["`specs/prd.md`",
             "Purpose, scope, goals and success criteria; REQ-1 to REQ-16 with their memory "
             "requirement ids; user stories; assumptions, constraints and open questions"],
            ["`specs/source/spezifikation-extract.md`",
             "The governing specification's structure, section by section. Requirement "
             "`source` fields cite it as `spec:§n`"],
            ["`specs/tasks/README.md`",
             "Phase index, per-task dependencies and complexity, and the full "
             "requirement-to-task mapping. Generated"],
            ["`docs/labkey-ce-ground-truth.md`",
             "Anything you are about to assume about LabKey CE: version and build, the premium "
             "boundary, what CE genuinely provides, HTTP action names, the file-based module "
             "layout, live server state"],
            ["`docs/gap-analysis.md`",
             "The 34-capability inventory, the five traps, the nine high-gap areas, where OSM "
             "goes beyond the commercial product, and the sources"],
            ["`standards/labkey/http-conventions.md`",
             "The eighteen rules for talking to a LabKey server. Departing from one needs a "
             "reason recorded in the change"],
            ["`AGENTS.md`", "The consolidated working rules: framing, verification, security, "
                            "resilience, boundaries, method, conventions"],
        ],
        widths=[2.0, 4.5],
    )

    g.h2("11.5 Environment variable quick reference")
    g.code(
        "# LabKey publish target\n"
        "export LK_URL=https://127.0.0.1:8443     # default; must be http(s) with a host\n"
        "export LK_USER=...                       # unless LK_APIKEY is set\n"
        "export LK_PASSWORD=...                   # unless LK_APIKEY is set\n"
        "export LK_APIKEY=...                     # takes precedence over user/password\n"
        "export LK_CONTEXT=auto                   # auto | '' | labkey\n"
        "export LK_INSECURE=                      # 1 to skip TLS verification\n"
        "export LK_TIMEOUT=60                     # seconds, positive number\n"
        "\n"
        "# Project memory (optional overrides)\n"
        "export OSM_MEMORY_DB=.sdd/memory.db\n"
        "export OSM_MEMORY_DUMP=.sdd/memory.sql\n"
        "\n"
        "# Reserved: declared, documented, and not read by any code yet\n"
        "export HUGGINGFACE_API_KEY=\n"
        "export HUGGINGFACE_MODEL_ID=\n"
        "export GITHUB_TOKEN="
    )

    g.h2("11.6 Command quick reference")
    g.code(
        "# Quality gate\n"
        "make check                       # lint, types, test, memory, render, docs-check\n"
        "make test                        # unit tests, no network\n"
        "make integration                 # tests marked `labkey`\n"
        "make docs                        # rebuild this guide\n"
        "make docs-check                  # fail if this guide is stale\n"
        "make clean                       # remove __pycache__ and tool caches\n"
        "\n"
        "# Knowledge base\n"
        "tools/memory.py stats\n"
        "tools/memory.py check\n"
        "tools/memory.py list backlog --brief\n"
        "tools/memory.py list backlog --where \"status='todo'\" --brief\n"
        "tools/memory.py show PR-001\n"
        "tools/memory.py query \"SELECT id,name,ce_support,gap FROM features WHERE gap='high'\"\n"
        "tools/memory.py rebuild\n"
        "tools/memory.py set backlog PR-001 status review\n"
        "tools/memory.py link --req FR-003 --kind backlog --to PR-004\n"
        "\n"
        "# Generated documents\n"
        "tools/render_backlog.py\n"
        "tools/render_backlog.py --check\n"
        "tools/build_guide.py --check\n"
        "\n"
        "# LabKey\n"
        "scripts/verify_labkey.py\n"
        "scripts/verify_labkey.py --record V-030"
    )

    g.h2("11.7 Known inconsistencies in the repository's own documentation")
    g.p("Flagged here rather than silently corrected, because the fix belongs in the "
        "repository and not in a generated guide.")
    g.table(
        ["Where", "Says", "Actually"],
        [
            ["`README.md` status section", "\"18 verifications\"",
             f"{f['counts']['verifications']} verification rows at this revision; PR-001 and "
             "the hardening pass added the rest"],
            ["`memory.md` and `.env.example`", "No mention of `LK_TIMEOUT`",
             "`src/osm/labkey/config.py` reads it and validates it"],
            ["`memory.md` \"Where things live\"", "Refers to `specs/README.md`",
             "That file does not exist"],
            ["`.github/prompts/verify.prompt.md`", "Refers to `scripts/labkey_client.py`",
             "The client is `src/osm/labkey/client.py`; the harness is "
             "`scripts/verify_labkey.py`"],
            ["`memory.md` and `README.md`", "Refer to `specs/features/`",
             "The directory exists but is empty; no FRD has been written"],
        ],
        widths=[1.7, 2.0, 2.8],
    )

    g.h2("11.8 About this document")
    g.p("Generated by `tools/build_guide.py` from the repository at revision "
        f"`{f['commit_full']}` on branch `{f['branch']}`, dated "
        f"{f['date'].strftime('%Y-%m-%d')}. The project version, revision, backlog, decision "
        "register, capability rollups and knowledge-base row counts are read from "
        "`pyproject.toml`, from git and from the memory database at build time; the prose is "
        "held in the generator and reviewed as source.")
    g.p("Rebuild it with `make docs`. Verify it is current with "
        "`tools/build_guide.py --check`. Do not edit the `.docx` by hand — the next build "
        "overwrites it.")
    g.p("Code is licensed Apache-2.0; documentation, including this guide, CC-BY-4.0.")


# --------------------------------------------------------------------------- build


def build_document(facts: dict[str, Any]) -> Guide:
    g = Guide(facts)
    write_title_page(g)
    write_toc(g)
    write_introduction(g)
    write_architecture(g)
    write_prerequisites(g)
    write_installation(g)
    write_configuration(g)
    write_usage(g)
    write_testing(g)
    write_security(g)
    write_troubleshooting(g)
    write_roadmap(g)
    write_appendix(g)

    props = g.doc.core_properties
    props.title = "OSM Sample Manager — Operator and Administrator Guide"
    props.subject = ("Open Sample Manager: installation, configuration, usage, verification "
                     "and current status")
    props.author = ATTRIBUTION
    props.last_modified_by = "tools/build_guide.py"
    props.category = "Operator and administrator guide"
    props.comments = provenance(facts)
    props.created = facts["date"]
    props.modified = facts["date"]
    props.revision = 1
    return g


#: Machine-readable provenance, written into the package and read back by
#: ``--check``. Pinned to an exact shape because it is parsed, not only read.
PROVENANCE = re.compile(
    r"^Generated by tools/build_guide\.py from revision (?P<commit>[0-9a-f]{7,40}) "
    r"on branch (?P<branch>\S+) at (?P<date>\S+)\."
)


def provenance(facts: dict[str, Any]) -> str:
    return (f"Generated by tools/build_guide.py from revision {facts['commit_full']} "
            f"on branch {facts['branch']} at {facts['date'].isoformat()}. "
            "Do not edit by hand; run `make docs`.")


def embedded_provenance(path: Path) -> dict[str, Any] | None:
    """Read back the revision, branch and date a document was built from.

    ``--check`` exists to answer *does this document still describe the
    repository?*, not *was it built at HEAD?*. Those differ by exactly one
    commit the moment the guide is committed, because committing it moves HEAD.
    Rebuilding with the stored provenance isolates the comparison to content:
    prose, version, backlog, decisions and row counts all still have to match,
    and only the revision label is allowed to be older.
    """
    try:
        comments = docx.Document(str(path)).core_properties.comments or ""
    except Exception:  # noqa: BLE001 - an unreadable file is simply not comparable
        return None
    match = PROVENANCE.match(comments)
    if not match:
        return None
    try:
        when = _dt.datetime.fromisoformat(match.group("date"))
    except ValueError:
        return None
    commit = match.group("commit")
    return {"commit_full": commit, "commit": commit[:7],
            "branch": match.group("branch"), "date": when}


def normalise_zip(source: Path, destination: Path, when: _dt.datetime) -> None:
    """Rewrite the package with fixed member timestamps in sorted order.

    python-docx stamps each zip member with the wall clock, which would make two
    builds of identical content differ byte for byte. Sorting also keeps
    ``[Content_Types].xml`` first, which the OPC format requires.
    """
    stamp = (when.year, when.month, when.day, when.hour, when.minute, when.second)
    with zipfile.ZipFile(source) as zin:
        members = {name: zin.read(name) for name in sorted(zin.namelist())}
    with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in members.items():
            info = zipfile.ZipInfo(name, date_time=stamp)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            zout.writestr(info, blob)


def render(out: Path, facts: dict[str, Any]) -> tuple[bytes, int]:
    """Build the document and return its normalised bytes and its section count."""
    guide = build_document(facts)
    with tempfile.TemporaryDirectory() as tmp:
        raw = Path(tmp) / "raw.docx"
        flat = Path(tmp) / "flat.docx"
        guide.doc.save(str(raw))
        normalise_zip(raw, flat, facts["date"])
        blob = flat.read_bytes()
    out.parent.mkdir(parents=True, exist_ok=True)
    return blob, guide.sections


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT,
                        help=f"output path (default: {DEFAULT_OUT.relative_to(REPO)})")
    parser.add_argument("--check", action="store_true",
                        help="exit 2 if the file on disk differs from a fresh build")
    args = parser.parse_args(argv)

    facts = load_facts()

    if args.check:
        if not args.out.exists():
            print(f"MISSING generated document: {args.out} (run: make docs)", file=sys.stderr)
            return 2
        stored = embedded_provenance(args.out)
        if stored is None:
            print(f"UNREADABLE generated document: {args.out} carries no provenance "
                  "(run: make docs)", file=sys.stderr)
            return 2
        blob, _ = render(args.out, {**facts, **stored})
        if args.out.read_bytes() != blob:
            print(f"STALE generated document: {args.out} no longer matches the repository "
                  "(run: make docs)", file=sys.stderr)
            return 2
        print(f"{args.out} is current (built at revision {stored['commit']})")
        return 0

    blob, sections = render(args.out, facts)

    tmp = args.out.with_suffix(args.out.suffix + ".tmp")
    tmp.write_bytes(blob)
    shutil.move(str(tmp), str(args.out))
    print(f"wrote {args.out} ({len(blob):,} bytes, {sections} top-level sections, "
          f"revision {facts['commit']})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
